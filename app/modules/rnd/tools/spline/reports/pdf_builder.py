from __future__ import annotations

import io
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Mapping

from docx.shared import RGBColor

# Shared templates folder: backend/app/utils/templates/
_TEMPLATES_DIR = Path(__file__).resolve().parents[5] / "utils" / "templates"


# ── DOCX ──────────────────────────────────────────────────────────────────────

def create_spline_docx(data: Mapping[str, Any], result: Mapping[str, Any]) -> io.BytesIO:
    from docx import Document

    doc = Document()

    # Title block
    title = doc.add_heading(
        "Calculation of Spline Parameters for Gear Box\nshaft and Half Gear Coupling", level=0
    )
    title.alignment = 1

    meta = doc.add_paragraph()
    meta.alignment = 1
    meta.add_run(f"Doc NO. - {data.get('doc_no') or 'PEW57-003-00'}\n")
    meta.add_run(f"Made By: {data.get('made_by') or ''}\n")
    meta.add_run(f"Checked By: {data.get('checked_by') or ''}\n")
    meta.add_run(f"Approved By: {data.get('approved_by') or ''}\n")
    meta.add_run(f"{data.get('doc_date') or ''}\n")
    meta.add_run(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 1. Introduction
    doc.add_heading("1  Introduction", level=1)
    doc.add_paragraph(
        "This document outlines the calculations involved in determining key parameters of "
        "splines used to transmit torque in mechanical systems. These parameters include the "
        "pitch diameter, base diameter, tooth thickness, fillet radius, torque capacity, and shear stress."
    )

    # 2. Spline Specifications
    doc.add_heading("2  Spline Specifications", level=1)
    for item in [
        "Number of teeth (Z)",
        "Diametral pitch (P) (teeth per unit diameter)",
        "Pressure angle (ϕ)",
        "Material properties",
    ]:
        doc.add_paragraph(f"  •  {item}")

    # 3. Calculations
    doc.add_heading("3  Calculations", level=1)
    doc.add_heading("3.1  Pitch Diameter (D)", level=2)
    doc.add_paragraph("D = Z / P")
    doc.add_heading("3.2  Base Diameter (Db)", level=2)
    doc.add_paragraph("Db = D × cos(ϕ)")
    doc.add_heading("3.3  Tooth Thickness (t)", level=2)
    doc.add_paragraph("t = π / (2P)")
    doc.add_heading("3.4  Fillet Radius (r)", level=2)
    doc.add_paragraph("Determined per spline standard.")
    doc.add_heading("3.5  Torque Capacity (T)", level=2)
    doc.add_paragraph("T = (D × H × L × Z × τ_a) / 2")
    doc.add_heading("3.6  Shear Stress (τ)", level=2)
    doc.add_paragraph("τ = (2 × T) / (D × L × Z × H)")

    # 4. Example
    doc.add_heading("4  Example", level=1)

    D = result.get('D', 0)
    Z = result.get('Z', 0)
    P = result.get('P', 0)
    H = result.get('H', 0)
    phi = data.get('pressure_angle') or data.get('pressure_angle_deg', 0)
    L = data.get('length_engagement', 0)
    mat = data.get('material_type', 'given material')
    Syt = result.get('yield_strength', 0)
    tau_a = result.get('allowable_shear', 0)
    T_cap = result.get('torque_capacity', 0)
    t_th = result.get('tooth_thickness', 0)
    base_d = result.get('base_diameter', 0)
    W = data.get('loco_weight', 0)
    n_axl = data.get('number_axles', 0)
    n_wpa = data.get('wheels_per_axle', 0)
    spd = data.get('speed', 0)
    w_dia = data.get('wheel_diameter', 0)
    mu = data.get('friction_coeff', 0)
    A = result.get('A', 0)
    C = result.get('C', 0)
    roll = result.get('rolling_resistance_n', 0) / 1000.0
    start = result.get('starting_resistance_n', 0) / 1000.0
    total = result.get('total_resistance_n', 0) / 1000.0
    T1 = result.get('working_torque_wheel', 0)
    tau = result.get('shear_stress_wheel', 0)
    fos = result.get('safety_factor', 0)

    # 4.0.1 Pitch Diameter
    doc.add_heading("4.0.1  Pitch Diameter (D)", level=2)
    p = doc.add_paragraph(f"D = Z / P = {Z:.2f} / {P:.2f} = {D:.2f} mm")

    # 4.0.2 Base Diameter
    doc.add_heading("4.0.2  Base Diameter (Db)", level=2)
    p = doc.add_paragraph()
    p.add_run(f"ϕ = {phi}°\n")
    p.add_run(f"Db = {D:.2f} × cos({phi}) = {base_d:.2f} mm")

    # 4.0.3 Tooth Thickness
    doc.add_heading("4.0.3  Tooth Thickness (t)", level=2)
    p = doc.add_paragraph()
    p.add_run(f"P = Z / D = {int(Z)} / {D:.2f} = {P:.2f} teeth/mm\n")
    p.add_run(f"t = π / (2P) = π / (2 × {P:.2f}) = {t_th:.2f} mm")

    # 4.0.4 Torque Capacity
    doc.add_heading("4.0.4  Torque Capacity (T)", level=2)
    p = doc.add_paragraph()
    p.add_run(f"D={D:.2f} mm,  H=(OD−ID)/2={H:.2f} mm,  L={L} mm,  Z={int(Z)}\n")
    p.add_run(f"τ_a = 0.577 × Syt = 0.577 × {Syt:.2f} = {tau_a:.2f} MPa\n")
    p.add_run(f"T = ({D:.2f} × {H:.2f} × {L} × {int(Z)} × {tau_a:.2f}) / 2\n")
    p.add_run(f"T = {T_cap:.2f} kNm").bold = True

    # 4.0.5 Working Torque per Wheel
    doc.add_heading("4.0.5  Working Torque per Wheel (T1)", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Loco Weight: {W} t,  Axles: {n_axl},  Wheels/Axle: {n_wpa}\n")
    p.add_run(f"Speed: {spd} km/h,  Wheel Diameter: {w_dia} m,  μ: {mu}\n")
    p.add_run(f"A = 0.647 + 13.17/(LocoWt/Axles) = {A:.4f}\n")
    p.add_run(f"B = 0.00933,  C = 0.057/LocoWt = {C:.6f}\n")
    p.add_run(f"Rolling Resistance = {roll:.2f} kN\n").bold = True
    p.add_run(f"Starting Resistance = {start:.2f} kN\n").bold = True
    p.add_run(f"Total Resistance = {total:.2f} kN  (gradient & curvature = 0)\n")
    p.add_run(f"T1 = Total Resistance × 1000 × Wheel Radius / No. Wheels = {T1:.2f} Nm\n").bold = True

    # 4.0.6 Working Shear Stress
    doc.add_heading("4.0.6  Working Shear Stress (τ)", level=2)
    p = doc.add_paragraph()
    p.add_run(f"τ = (2 × T1 × 10³) / (D × L × Z × H)\n")
    p.add_run(f"  = (2 × {T1:.2f} × 10³) / ({D:.2f} × {L} × {int(Z)} × {H:.2f})\n")
    p.add_run(f"τ = {tau:.2f} MPa\n").bold = True
    p.add_run(f"FOS = τ_a / τ = {tau_a:.2f} / {tau:.2f} = {fos:.2f}\n")
    try:
        r = p.add_run(f"FOS = {fos:.2f}")
        r.bold = True
        r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    except Exception:
        p.add_run(f"FOS = {fos:.2f}")

    # 4.0.7 Summary
    doc.add_heading("4.0.7  Summary of Results", level=2)
    verdict = result.get('verdict', '')
    for label, value in [
        ("Tooth Thickness (t)", f"{t_th:.2f} mm"),
        ("Torque Capacity (T)", f"{T_cap:.2f} kNm"),
        ("Working Torque per Wheel (T1)", f"{T1:.2f} Nm"),
        ("Working Shear Stress (τ)", f"{tau:.2f} MPa"),
        ("Allowable Shear Stress (τ_a)", f"{tau_a:.2f} MPa"),
        ("FOS at working torque", f"{fos:.2f}"),
        ("Verdict", verdict),
    ]:
        doc.add_paragraph(f"  •  {label}: {value}")

    # 5. Conclusion
    doc.add_heading("5  Conclusion", level=1)
    doc.add_paragraph(
        "This document provides the theoretical and practical calculations necessary for "
        "designing splines capable of transmitting torque effectively in mechanical systems."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── PDF (LaTeX) ───────────────────────────────────────────────────────────────

def generate_spline_pdf(data: Mapping[str, Any], result: Mapping[str, Any]) -> io.BytesIO:
    try:
        from jinja2 import Environment, FileSystemLoader
    except ImportError:
        raise Exception("Jinja2 not installed. Run: pip install Jinja2")

    template_name = "spline_template.tex"
    template_path = _TEMPLATES_DIR / template_name
    if not template_path.exists():
        raise Exception(f"LaTeX template not found: {template_path}")

    env = Environment(loader=FileSystemLoader(str(_TEMPLATES_DIR)))
    template = env.get_template(template_name)
    context = {"data": dict(data), "result": dict(result)}
    latex_content = template.render(**context)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        tex_file = tmp_path / "spline_report.tex"
        tex_file.write_text(latex_content, encoding="utf-8")

        for asset in ["logo.JPG", "logo-1.JPG"]:
            src = _TEMPLATES_DIR / asset
            if src.exists():
                shutil.copy(src, tmp_path / asset)

        pdf_path = tmp_path / "spline_report.pdf"
        last_res = None
        for compiler in ["pdflatex", "xelatex", "lualatex"]:
            try:
                last_res = subprocess.run(
                    [compiler, "-interaction=nonstopmode", "spline_report.tex"],
                    cwd=tmp, capture_output=True, text=True, timeout=30,
                )
            except FileNotFoundError:
                last_res = None
                continue
            if pdf_path.exists():
                return io.BytesIO(pdf_path.read_bytes())

        diag = []
        if last_res:
            diag.append(f"returncode={last_res.returncode}")
            diag.append(last_res.stdout or "")
            diag.append(last_res.stderr or "")
        log = tmp_path / "spline_report.log"
        if log.exists():
            try:
                diag.append(log.read_text(encoding="utf-8", errors="ignore")[-2000:])
            except Exception:
                pass
        raise Exception("LaTeX compilation failed.\n" + "\n".join(diag))
