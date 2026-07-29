"""DOCX report builder for Spline tool — mirrors the LaTeX PDF structure exactly."""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Mapping

try:
    from docx import Document
    from docx.shared import RGBColor, Pt
except ImportError:
    Document = None
    RGBColor = None
    Pt = None


def _bold(p: Any, text: str) -> None:
    p.add_run(text).bold = True


def create_spline_docx(data: Mapping[str, Any], result: Mapping[str, Any]) -> io.BytesIO:
    """Generate a DOCX report matching the spline PDF structure.

    Parameters
    ----------
    data:   SplineInput dict (raw user inputs)
    result: dict returned by calculate_spline_mode()
    """
    if Document is None:
        raise Exception("python-docx library is not installed. Please run: pip install python-docx")

    doc = Document()

    # ── Title block ───────────────────────────────────────────────────────────
    title = doc.add_heading("Calculation of Spline Parameters for Gear Box\nshaft and Half Gear Coupling", level=0)
    title.alignment = 1  # center

    meta = doc.add_paragraph()
    meta.alignment = 1
    meta.add_run(f"Doc NO. - {data.get('doc_no', 'PEW57-003-00')}\n")
    meta.add_run(f"Made By: {data.get('made_by', 'Gauransh Kaushik')}\n")
    meta.add_run(f"Checked By: {data.get('checked_by', 'Jasbir Singh')}\n")
    meta.add_run(f"Approved By: {data.get('approved_by', 'Madhav Arora')}\n")
    meta.add_run(f"{data.get('doc_date', 'November 19, 2024')}\n")
    meta.add_run(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # ── 1. Introduction ───────────────────────────────────────────────────────
    doc.add_heading("1  Introduction", level=1)
    doc.add_paragraph(
        "This document outlines the calculations involved in determining key parameters of "
        "splines used to transmit torque in mechanical systems. These parameters include the "
        "pitch diameter, base diameter, tooth thickness, fillet radius, torque capacity, and shear stress."
    )

    # ── 2. Spline Specifications ──────────────────────────────────────────────
    doc.add_heading("2  Spline Specifications", level=1)
    for item in [
        "Number of teeth (Z)",
        "Diametral pitch (P) (teeth per unit diameter)",
        "Pressure angle (\u03d5)",
        "Material properties",
    ]:
        doc.add_paragraph(f"  \u2022  {item}")

    # ── 3. Calculations ───────────────────────────────────────────────────────
    doc.add_heading("3  Calculations", level=1)

    doc.add_heading("3.1  Pitch Diameter (D)", level=2)
    doc.add_paragraph("The pitch diameter is calculated using the formula:")
    doc.add_paragraph("  D = Z / P")
    doc.add_paragraph("where Z is the number of teeth and P is the diametral pitch.")

    doc.add_heading("3.2  Base Diameter (Db)", level=2)
    doc.add_paragraph("The base diameter, necessary for the involute profile, is calculated as:")
    doc.add_paragraph("  Db = D \u00d7 cos(\u03d5)")

    doc.add_heading("3.3  Tooth Thickness (t)", level=2)
    doc.add_paragraph("Tooth thickness at the pitch diameter:")
    doc.add_paragraph("  t = \u03c0 / (2P)")

    doc.add_heading("3.4  Fillet Radius (r)", level=2)
    doc.add_paragraph(
        "The fillet radius is determined based on the spline standard and ensures the "
        "teeth are not prone to breakage."
    )

    doc.add_heading("3.5  Torque Capacity (T)", level=2)
    doc.add_paragraph("Calculated by:")
    doc.add_paragraph("  T = (D \u00d7 H \u00d7 L \u00d7 Z \u00d7 \u03c4_a) / 2")

    doc.add_heading("3.6  Shear Stress (\u03c4)", level=2)
    doc.add_paragraph("Shear stress in the spline teeth:")
    doc.add_paragraph("  \u03c4 = (2 \u00d7 T) / (D \u00d7 L \u00d7 Z \u00d7 H)")

    # ── 4. Example ────────────────────────────────────────────────────────────
    doc.add_heading("4  Example", level=1)
    doc.add_paragraph(
        "To illustrate the calculation, consider the input values used in this report. "
        "The following results are obtained by applying the formulas of Section 3 to the measured data."
    )

    # Shortcuts
    D   = result.get('D', 0)
    Z   = result.get('Z', 0)
    P   = result.get('P', 0)
    H   = result.get('H', 0)
    phi = data.get('pressure_angle', data.get('pressure_angle_deg', 0))
    L   = data.get('length_engagement', 0)
    mat = data.get('material_type', 'given material')
    Syt = result.get('yield_strength', 0)
    tau_a = result.get('allowable_shear', 0)
    T_cap = result.get('torque_capacity', 0)
    t_th  = result.get('tooth_thickness', 0)
    base_d = result.get('base_diameter', 0)

    W     = data.get('loco_weight', 0)
    n_axl = data.get('number_axles', 0)
    n_wpa = data.get('wheels_per_axle', 0)
    spd   = data.get('speed', 0)
    w_dia = data.get('wheel_diameter', 0)
    mu    = data.get('friction_coeff', 0)
    A     = result.get('A', 0)
    C     = result.get('C', 0)
    roll  = result.get('rolling_resistance_n', 0) / 1000.0
    start = result.get('starting_resistance_n', 0) / 1000.0
    total = result.get('total_resistance_n', 0) / 1000.0
    T1    = result.get('working_torque_wheel', 0)
    tau   = result.get('shear_stress_wheel', 0)
    fos   = result.get('safety_factor', 0)

    # 4.0.1 Pitch Diameter
    doc.add_heading("4.0.1  Pitch Diameter (D)", level=2)
    p = doc.add_paragraph(f"Given that the pitch diameter D is:\n")
    p.add_run(f"  D = Z / P = {Z:.2f} / {P:.2f} = {D:.2f} mm")

    # 4.0.2 Base Diameter
    doc.add_heading("4.0.2  Base Diameter (Db)", level=2)
    p = doc.add_paragraph("To calculate the base diameter Db, we need the pressure angle \u03d5.\n")
    p.add_run(f"  Db = D \u00d7 cos(\u03d5)\n")
    p.add_run(f"  \u03d5 = {phi}\u00b0\n")
    p.add_run(f"  Db = {D:.2f} \u00d7 cos({phi}) = {base_d:.2f} mm")

    # 4.0.3 Tooth Thickness
    doc.add_heading("4.0.3  Tooth Thickness (t)", level=2)
    p = doc.add_paragraph("To calculate the tooth thickness at the pitch diameter, we need the diametral pitch P:\n")
    p.add_run(f"  P = Z / D = {int(Z)} / {D:.2f} = {P:.2f} teeth/mm\n")
    p.add_run(f"  t = \u03c0 / (2P) = \u03c0 / (2 \u00d7 {P:.2f}) = {t_th:.2f} mm")

    # 4.0.4 Torque Capacity
    doc.add_heading("4.0.4  Torque Capacity (T)", level=2)
    p = doc.add_paragraph("The torque capacity is calculated using the formula:\n")
    p.add_run("  T = (D \u00d7 H \u00d7 L \u00d7 Z \u00d7 \u03c4_a) / 2\n\n")
    p.add_run("Where:\n")
    p.add_run(f"  D = {D:.2f} mm\n")
    p.add_run(f"  H = (OD \u2212 ID) / 2 = {H:.2f} mm\n")
    p.add_run(f"  L = {L} mm\n")
    p.add_run(f"  Z = {int(Z)}\n")
    p.add_run(f"  For {mat}:\n")
    p.add_run(f"    \u03c4_a = 0.577 \u00d7 Syt = 0.577 \u00d7 {Syt:.2f} MPa = {tau_a:.2f} MPa\n\n")
    p.add_run(f"  T = ({D:.2f} \u00d7 {H:.2f} \u00d7 {L} \u00d7 {int(Z)} \u00d7 {tau_a:.2f}) / 2\n")
    r = p.add_run(f"  T = {T_cap:.2f} kNm")
    r.bold = True

    # 4.0.5 Working Torque per Wheel
    doc.add_heading("4.0.5  Working Torque per Wheel (T1)", level=2)
    p = doc.add_paragraph("Working conditions:\n")
    p.add_run(f"  Locomotive Weight: {W} tons\n")
    p.add_run(f"  Number of Axles: {n_axl}\n")
    p.add_run(f"  Number of wheels per axle: {n_wpa}\n")
    p.add_run(f"  Speed: {spd} km/h\n")
    p.add_run(f"  Wheel Diameter: {w_dia} m\n")
    p.add_run(f"  Coefficient of Friction: {mu}\n")

    p = doc.add_paragraph("Rolling Resistance = (A + B \u00d7 Speed + C \u00d7 Speed\u00b2) \u00d7 Loco. Weight \u00d7 g\n")
    p.add_run(f"  A = 0.647 + (13.17 / (LocoWt / Axles)) = 0.647 + (13.17 / ({W} / {n_axl})) = {A:.2f}\n")
    p.add_run(f"  B = 0.00933\n")
    p.add_run(f"  C = 0.057 / LocoWt = 0.057 / {W} = {C:.2f}\n")
    r = p.add_run(f"  Rolling Resistance = {roll:.2f} kN")
    r.bold = True

    p = doc.add_paragraph(f"Starting Resistance = 6 \u00d7 Loco. Weight \u00d7 g\n")
    r = p.add_run(f"  Starting Resistance = {start:.2f} kN")
    r.bold = True

    p = doc.add_paragraph(
        "Total Resistance = Rolling Resistance + Gradient Resistance + "
        "Curvature Resistance + Starting Resistance\n"
    )
    p.add_run(f"  Total Resistance = {total:.2f} kN  "
              f"(\u2235 Gradient Resistance and Curvature Resistance are taken as zero)\n\n")
    p.add_run(f"  T1 = Total Resistance \u00d7 1000 \u00d7 Wheel Radius / Number of Wheels\n")
    r = p.add_run(f"  T1 = {T1:.2f} Nm")
    r.bold = True

    # 4.0.6 Working Shear Stress
    doc.add_heading("4.0.6  Working Shear Stress (\u03c4)", level=2)
    p = doc.add_paragraph("For the shear stress calculation, the formula is:\n")
    p.add_run("  \u03c4 = (2 \u00d7 T1) / (D \u00d7 L \u00d7 Z \u00d7 H)\n\n")
    p.add_run("Where:\n")
    p.add_run(f"  D = {D:.2f} mm\n")
    p.add_run(f"  H = {H:.2f} mm\n")
    p.add_run(f"  L = {L} mm\n")
    p.add_run(f"  Z = {int(Z)}\n")
    p.add_run(f"  T1 = {T1:.2f} Nm\n\n")
    p.add_run("Substituting:\n")
    p.add_run(
        f"  \u03c4 = (2 \u00d7 {T1:.2f} \u00d7 10\u00b3) / "
        f"({D:.2f} \u00d7 {L} \u00d7 {int(Z)} \u00d7 {H:.2f})\n"
    )
    r = p.add_run(f"  \u03c4 = {tau:.2f} MPa")
    r.bold = True

    p = doc.add_paragraph("Factor of Safety (FOS):\n")
    p.add_run(f"  FOS = \u03c4_a / \u03c4 = {tau_a:.2f} MPa / {tau:.2f} MPa\n")
    r = p.add_run(f"  FOS = {fos:.2f}")
    r.bold = True
    try:
        r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    except Exception:
        pass

    # 4.0.7 Summary
    doc.add_heading("4.0.7  Summary of Results", level=2)
    summary = [
        ("Tooth Thickness (t)",           f"{t_th:.2f} mm"),
        ("Torque Capacity (T)",           f"{T_cap:.2f} kNm"),
        ("Working Torque per Wheel (T1)", f"{T1:.2f} Nm"),
        ("Working Shear Stress (\u03c4)",  f"{tau:.2f} MPa"),
        ("Allowable Shear Stress (\u03c4_a)", f"{tau_a:.2f} MPa"),
        ("FOS at working torque",         f"{fos:.2f}"),
    ]
    for label, value in summary:
        doc.add_paragraph(f"  \u2022  {label}: {value}")

    # ── 5. Conclusion ─────────────────────────────────────────────────────────
    doc.add_heading("5  Conclusion", level=1)
    doc.add_paragraph(
        "This document provides the theoretical and practical calculations necessary for "
        "designing splines capable of transmitting torque effectively in mechanical systems."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
