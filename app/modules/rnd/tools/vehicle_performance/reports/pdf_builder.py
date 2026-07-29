import io
import math
from collections import defaultdict
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import numpy as np

plt: Any = None
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt  # type: ignore
    import matplotlib.ticker as mticker  # type: ignore
except Exception:
    plt = None

_G = 9.81


# ── Resistance formulas (exact copy from service.py) ────────────────────────

def _rr_loco(speed_kmh: float, weight_ton: float, num_axles: int) -> float:
    if weight_ton <= 0 or num_axles <= 0:
        return 0.0
    A = 0.647 + (13.17 / (weight_ton / num_axles))
    B = 0.00933
    C = (0.057 / weight_ton) if weight_ton > 0 else 0
    return (A + B * speed_kmh + C * speed_kmh ** 2) * weight_ton * _G


def _rr_wagon(speed_kmh: float, weight_ton: float) -> float:
    if weight_ton <= 0:
        return 0.0
    return (0.6438797 + 0.01047218 * speed_kmh + 0.00007323 * speed_kmh ** 2) * weight_ton * _G


def _grad_res(weight_ton: float, slope_pct: float) -> float:
    return weight_ton * 1000 * _G * slope_pct / 100.0


def _curve_res(weight_ton: float, curve_deg: float) -> float:
    return 0.4 * weight_ton * curve_deg * _G


def _start_loco(weight_ton: float) -> float:
    return 6.0 * weight_ton * _G


def _start_wagon(weight_ton: float) -> float:
    return 4.0 * weight_ton * _G


def _interp_torque(engine_rpm: float, torque_curve: Dict[int, float]) -> float:
    if not torque_curve:
        return 0.0
    rpms = sorted(torque_curve.keys())
    torques = [torque_curve[r] for r in rpms]
    return float(np.interp(engine_rpm, rpms, torques))


# ── Table helpers ────────────────────────────────────────────────────────────

def _hdr_row(table, *texts):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        # Remove \n from headers — use space instead
        cell.text = str(text).replace('\n', ' ')
        runs = cell.paragraphs[0].runs
        if not runs:
            continue
        run = runs[0]
        run.bold = True
        run.font.size = Pt(9)
        # dark header background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E293B')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _add_row(table, *texts):
    from docx.shared import Pt
    rc = table.add_row().cells
    for i, text in enumerate(texts):
        rc[i].text = str(text)
        runs = rc[i].paragraphs[0].runs
        if runs:
            runs[0].font.size = Pt(9)
    return rc


def _shade_row(row_cells, hex_fill: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in row_cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_fill)
        tcPr.append(shd)


def _set_col_widths(table, widths_cm: List[float]):
    # 1 cm ≈ 567 twips (dxa units used by Word)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), str(int(widths_cm[i] * 567)))
                tcW.set(qn('w:type'), 'dxa')


# ── Main builder ─────────────────────────────────────────────────────────────

def create_vehicle_performance_docx_report(inputs: Dict[str, Any], results: Any) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if not isinstance(results, dict):
        raise RuntimeError(f"Expected results dict, got {type(results).__name__}")

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

    # ── Normal font ───────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ══════════════════════════════════════════════════════════════════
    # TITLE BLOCK
    # ══════════════════════════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('VEHICLE PERFORMANCE CALCULATION REPORT')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xF9, 0x73, 0x16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('Tractive Effort  ·  Shunting Capability  ·  Speed vs Slope Analysis').font.size = Pt(9)

    # generated timestamp line
    gen_p = doc.add_paragraph()
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_p.add_run(f'Generated: {datetime.now().strftime("%d %B %Y  %H:%M")}').font.size = Pt(9)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 1 — INPUT PARAMETERS
    # ══════════════════════════════════════════════════════════════════
    h = doc.add_heading('1.  Input Parameters', level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Extract all inputs
    gvw_kg     = float(inputs.get('loco_gvw_kg', 0))
    gvw_ton    = gvw_kg / 1000.0
    max_spd    = float(inputs.get('max_speed_kmh', 0))
    num_axles  = int(inputs.get('num_axles', 2))
    rear_axle  = float(inputs.get('rear_axle_ratio', 1.0))
    gear_ratios= inputs.get('gear_ratios', [1.0])
    shunt_t    = float(inputs.get('shunting_load_t', 0))
    peak_kw    = float(inputs.get('peak_power_kw', 0))
    mu         = float(inputs.get('friction_mu', 0.3))
    wheel_dia  = float(inputs.get('wheel_dia_m', 0.73))
    wheel_r    = wheel_dia / 2.0
    min_rpm    = int(inputs.get('min_rpm', 100))
    max_rpm    = int(inputs.get('max_rpm', 2500))
    max_slope  = float(inputs.get('max_slope', 3.5))
    slope_unit = inputs.get('slope_unit', '%')
    max_curve  = float(inputs.get('max_curve', 5.0))
    curve_unit = inputs.get('curve_unit', 'degree')
    torque_curve: Dict[int, float] = inputs.get('torque_curve', {})

    # Convert slope/curve for calculations
    if slope_unit == 'degree':
        max_slope_pct = math.tan(math.radians(max_slope)) * 100.0
    else:
        max_slope_pct = max_slope

    if curve_unit == 'm':
        curve_deg = (1750.0 / max_curve) if max_curve > 0 else 0.0
    else:
        curve_deg = max_curve

    inp_table = doc.add_table(rows=1, cols=4)
    inp_table.style = 'Table Grid'
    _hdr_row(inp_table, 'Parameter', 'Value', 'Parameter', 'Value')
    inp_rows = [
        ('GVW',                  f'{gvw_kg:,.0f} kg  ({gvw_ton:.2f} t)',
         'Peak Power',           f'{peak_kw:.0f} kW  ({peak_kw*1.341:.0f} HP)'),
        ('No. of Axles',         str(num_axles),
         'Friction Coeff. (μ)',  f'{mu:.3f}'),
        ('Max Speed',            f'{max_spd:.0f} km/h',
         'Wheel Diameter',       f'{wheel_dia:.3f} m'),
        ('Rear Axle Ratio',      f'{rear_axle:.3f}',
         'Min RPM',              str(min_rpm)),
        ('Gear Ratio(s)',         ', '.join(f'{g:.2f}' for g in gear_ratios),
         'Max RPM',              str(max_rpm)),
        ('Shunting Load',        f'{shunt_t:.1f} t',
         'Max Slope',            f'{max_slope:.2f} {slope_unit}'),
        ('Max Curve',            f'{max_curve:.1f} ({curve_unit})',
         'Slope (internal %)',   f'{max_slope_pct:.4f} %'),
        ('Curve (internal °)',   f'{curve_deg:.4f} °',
         'Wheel Radius',         f'{wheel_r:.4f} m'),
    ]
    for even, (l1, v1, l2, v2) in enumerate(inp_rows):
        rc = _add_row(inp_table, l1, v1, l2, v2)
        if even % 2 == 0:
            _shade_row(rc, 'F8FAFC')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 2 — TORQUE CURVE
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('2.  Torque Curve', level=1)

    if torque_curve:
        max_tq   = max(torque_curve.values())
        max_tq_r = max(torque_curve, key=torque_curve.get)

        info_p = doc.add_paragraph()
        info_p.add_run(f'  Peak Torque: ').bold = True
        info_p.add_run(f'{max_tq:.1f} N·m  at  {max_tq_r} RPM')

        tq_tbl = doc.add_table(rows=1, cols=3)
        tq_tbl.style = 'Table Grid'
        _hdr_row(tq_tbl, 'RPM', 'Torque (N·m)', 'Power at RPM (kW)')
        for i, (rpm, tq) in enumerate(sorted(torque_curve.items())):
            pwr = (rpm * tq * 2 * math.pi) / (60 * 1000)
            rc = _add_row(tq_tbl, str(rpm), f'{tq:.1f}', f'{pwr:.2f}')
            if i % 2 == 0:
                _shade_row(rc, 'F8FAFC')
        _set_col_widths(tq_tbl, [3.5, 4.0, 5.0])
    else:
        doc.add_paragraph('  No torque curve data provided.')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 3 — TRACTION ANALYSIS (step by step)
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('3.  Traction Analysis', level=1)

    traction = results.get('traction_snapshot', {})
    gen_n     = float(traction.get('max_traction_generated_n', 0))
    slip_n    = float(traction.get('max_traction_slipping_n', 0))
    status    = traction.get('result_message', '—')
    actual_n  = min(gen_n, slip_n)

    # ── 3.1 Max traction generated ──
    doc.add_heading('3.1  Step 1 — Maximum Traction Generated by Drivetrain', level=2)
    p = doc.add_paragraph()
    p.add_run('Formula:\n').bold = True
    p.add_run('  Maximum Traction Force  =  2 × (Peak Torque × Gear Ratio × Rear Axle Ratio)  ÷  Wheel Diameter\n\n')

    if torque_curve:
        p.add_run('  Max Torque (T_max)  =  ').bold = False
        p.add_run(f'{max_tq:.2f} N·m').bold = True
        p.add_run(f'  at {max_tq_r} RPM\n')

    p.add_run('  Max Gear Ratio      =  ').bold = False
    p.add_run(f'{max(gear_ratios):.3f}\n').bold = True
    p.add_run('  Rear Axle Ratio     =  ').bold = False
    p.add_run(f'{rear_axle:.3f}\n').bold = True
    p.add_run('  Wheel Diameter      =  ').bold = False
    p.add_run(f'{wheel_dia:.3f} m\n').bold = True

    p = doc.add_paragraph()
    if torque_curve:
        p.add_run(f'  Maximum Traction Force  =  2 × ({max_tq:.2f} N·m  ×  {max(gear_ratios):.3f}  ×  {rear_axle:.3f})  ÷  {wheel_dia:.3f} m\n')
    p.add_run('  ').bold = False
    run = p.add_run(f'F_traction  =  {gen_n:,.2f} N')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    # ── 3.2 No-slip limit ──
    doc.add_heading('3.2  Step 2 — No-Slip (Adhesion) Limit', level=2)
    p = doc.add_paragraph()
    p.add_run('Formula:\n').bold = True
    p.add_run('  No-Slip Limit  =  Gross Vehicle Weight (tonnes) × Friction Coefficient × 1000 × Gravity\n\n')
    p.add_run(f'  Gross Vehicle Weight   =  {gvw_ton:.3f} tonnes\n')
    p.add_run(f'  Friction Coefficient   =  {mu:.3f}\n')
    p.add_run(f'  Gravity (g)            =  {_G} m/s²\n\n')
    p.add_run(f'  No-Slip Limit  =  {gvw_ton:.3f}  ×  {mu:.3f}  ×  1000  ×  {_G}\n')
    run = p.add_run(f'  No-Slip Limit  =  {slip_n:,.2f} N')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x15, 0x80, 0x3D)

    # ── 3.3 Verdict ──
    doc.add_heading('3.3  Step 3 — Effective Traction & Verdict', level=2)
    p = doc.add_paragraph()
    p.add_run('  Effective Traction  =  min(Maximum Traction Force,  No-Slip Limit)\n')
    p.add_run(f'  Effective Traction  =  min({gen_n:,.2f} N,  {slip_n:,.2f} N)\n')
    run = p.add_run(f'  Effective Traction  =  {actual_n:,.2f} N\n\n')
    run.bold = True
    run.font.size = Pt(11)

    # Status badge row
    verdict_tbl = doc.add_table(rows=1, cols=3)
    verdict_tbl.style = 'Table Grid'
    cells = verdict_tbl.rows[0].cells
    cells[0].text = 'Generated (N)'
    cells[1].text = 'No-Slip Limit (N)'
    cells[2].text = 'Status'
    for c in cells:
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)

    rc = verdict_tbl.add_row().cells
    rc[0].text = f'{gen_n:,.2f}'
    rc[1].text = f'{slip_n:,.2f}'
    rc[2].text = status
    fill = 'FEE2E2' if 'slip' in status.lower() else 'DCFCE7'
    _shade_row([rc[2]], fill)
    for c in rc:
        c.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 4 — RESISTANCE BREAKDOWN
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('4.  Resistance Force Breakdown', level=1)

    p = doc.add_paragraph()
    p.add_run('Resistance formulas used (empirical, Davis equation):\n').bold = True
    p.add_run('  Rolling Res (loco)  = (A + B·v + C·v²) × W × g  ')
    p.add_run('where A = 0.647 + 13.17/(W/n_axles),  B = 0.00933,  C = 0.057/W\n')
    p.add_run('  Gradient Res        = W_kg × g × slope% / 100\n')
    p.add_run('  Curvature Res       = 0.4 × W_ton × curve_deg × g\n')
    p.add_run('  Starting Res (loco) = 6.0 × W_ton × g\n')
    p.add_run('  Starting Res (wagon)= 4.0 × W_ton × g\n\n')

    # Loco resistance at key speeds
    doc.add_heading('4.1  Locomotive Resistance vs Speed', level=2)
    speeds_chk = [0, 5, 10, 15, 20, 25, 30, 40, 50]
    res_tbl = doc.add_table(rows=1, cols=7)
    res_tbl.style = 'Table Grid'
    _hdr_row(res_tbl, 'Speed (km/h)', 'Rolling (N)', 'Gradient (N)',
             'Curvature (N)', 'Starting (N)', 'Total Running (N)', 'Total incl. Start (N)')
    for i, spd in enumerate(speeds_chk):
        rr = _rr_loco(spd, gvw_ton, num_axles)
        gr = _grad_res(gvw_ton, max_slope_pct)
        cr = _curve_res(gvw_ton, curve_deg)
        sr = _start_loco(gvw_ton)
        running = rr + gr + cr
        total   = running + sr
        rc = _add_row(res_tbl,
                      str(spd),
                      f'{rr:.1f}',
                      f'{gr:.1f}',
                      f'{cr:.1f}',
                      f'{sr:.1f}',
                      f'{running:.1f}',
                      f'{total:.1f}')
        if i % 2 == 0:
            _shade_row(rc, 'F8FAFC')

    doc.add_paragraph()

    # Wagon resistance (if shunting load > 0)
    if shunt_t > 0:
        doc.add_heading(f'4.2  Wagon Resistance vs Speed  (Load = {shunt_t:.1f} t)', level=2)
        wg_tbl = doc.add_table(rows=1, cols=6)
        wg_tbl.style = 'Table Grid'
        _hdr_row(wg_tbl, 'Speed (km/h)', 'Rolling (N)', 'Gradient (N)',
                 'Curvature (N)', 'Starting (N)', 'Total (N)')
        for i, spd in enumerate(speeds_chk):
            rr = _rr_wagon(spd, shunt_t)
            gr = _grad_res(shunt_t, max_slope_pct)
            cr = _curve_res(shunt_t, curve_deg)
            sr = _start_wagon(shunt_t)
            total = rr + gr + cr + sr
            rc = _add_row(wg_tbl, str(spd),
                          f'{rr:.1f}', f'{gr:.1f}', f'{cr:.1f}', f'{sr:.1f}', f'{total:.1f}')
            if i % 2 == 0:
                _shade_row(rc, 'F8FAFC')
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 5 — GEAR-WISE TRACTION AT REFERENCE SPEEDS
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('5.  Gear-wise Tractive Effort at Key Speeds', level=1)
    p = doc.add_paragraph()
    p.add_run('Formula per operating point:\n').bold = True
    p.add_run('  Wheel RPM          =  (Speed in m/s  ÷  (π × Wheel Diameter))  × 60\n')
    p.add_run('  Engine RPM         =  Wheel RPM  ×  Gear Ratio  ×  Rear Axle Ratio\n')
    p.add_run('  Torque (N·m)       =  Interpolate Engine RPM on Torque Curve  [limited to Peak Power]\n')
    p.add_run('  Wheel Force (N)    =  2 × (Torque × Gear Ratio × Rear Axle Ratio)  ÷  Wheel Diameter\n')
    p.add_run('  Actual Traction    =  min(Wheel Force,  No-Slip Limit)\n\n')

    ref_speeds = [5, 10, 15, 20, 25, 30]
    col_headers = ['Speed (km/h)'] + [f'Gear Ratio {g:.2f} — Actual Traction (N)' for g in gear_ratios] + ['No-Slip Limit (N)']
    gear_tbl = doc.add_table(rows=1, cols=len(col_headers))
    gear_tbl.style = 'Table Grid'
    _hdr_row(gear_tbl, *col_headers)

    for i, spd in enumerate(ref_speeds):
        spd_mps = spd / 3.6
        row_vals = [str(spd)]
        for gr in gear_ratios:
            w_rpm = (spd_mps / (math.pi * wheel_dia)) * 60 if wheel_dia > 0 else 0
            e_rpm = w_rpm * gr * rear_axle
            tq = _interp_torque(e_rpm, torque_curve) if torque_curve else 0.0
            if e_rpm > 0:
                pwr_check = (e_rpm * tq * 2 * math.pi) / (60 * 1000)
                if pwr_check > peak_kw:
                    tq = (peak_kw * 60 * 1000) / (e_rpm * 2 * math.pi)
            f_wheel = 2 * (tq * gr * rear_axle) / wheel_dia if wheel_dia > 0 else 0
            f_act = min(f_wheel, slip_n)
            row_vals.append(f'{f_act:,.1f}')
        row_vals.append(f'{slip_n:,.2f}')
        rc = _add_row(gear_tbl, *row_vals)
        if i % 2 == 0:
            _shade_row(rc, 'F8FAFC')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 6 — SPEED vs SLOPE TABLE
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('6.  Speed vs Slope  (Max Achievable Speed)', level=1)
    p = doc.add_paragraph()
    p.add_run(f'  Shunting Load: {shunt_t:.1f} t    '
              f'Max Slope: {max_slope:.2f} {slope_unit}    '
              f'Max Curve: {max_curve:.1f} ({curve_unit})\n')

    speed_slope: List[Dict] = results.get('speed_slope_table', [])
    if speed_slope:
        svs_tbl = doc.add_table(rows=1, cols=2)
        svs_tbl.style = 'Table Grid'
        _hdr_row(svs_tbl, 'Slope (%)', 'Max Achievable Speed (km/h)')
        for i, row in enumerate(speed_slope):
            sl  = row.get('slope', 0)
            spd = row.get('max_speed_kmh', 0)
            rc  = _add_row(svs_tbl, f'{sl:.2f}', f'{spd:.2f}')
            # Highlight zero-speed rows
            if spd == 0:
                _shade_row(rc, 'FEE2E2')
            elif i % 2 == 0:
                _shade_row(rc, 'F8FAFC')
        _set_col_widths(svs_tbl, [4.5, 6.0])
    else:
        doc.add_paragraph('  No speed vs slope data available.')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 7 — GRAPHS
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('7.  Performance Charts', level=1)

    plot_data: Dict[str, Any] = results.get('plot_data', {})

    if plt is None:
        doc.add_paragraph('  Matplotlib not installed — charts not included.')
    elif not isinstance(plot_data, dict):
        doc.add_paragraph('  No plot data available.')
    else:
        SLOPE_COLORS = [
            '#1D4ED8', '#059669', '#D97706', '#DC2626', '#7C3AED',
            '#0891B2', '#65A30D', '#EA580C', '#9333EA', '#0F766E',
        ]

        for chart_title, data_key, ylabel in [
            ('Tractive Effort vs Speed',       'tractive_effort_plot',    'Tractive Effort (N)'),
            ('Shunting Capability vs Speed',   'shunting_capability_plot','Shunting Capability (tons)'),
        ]:
            raw: List[Dict] = plot_data.get(data_key) or []
            if not raw:
                doc.add_paragraph(f'  No data for "{chart_title}".')
                continue

            # Group by (slope, gear)
            series: Dict[str, List[Tuple[float, float]]] = defaultdict(list)
            for pt in raw:
                spd = pt.get('speed_kmh')
                val = pt.get('value')
                sl  = pt.get('slope')
                g   = pt.get('gear')
                if spd is None or val is None:
                    continue
                key = f'Slope {sl}%  (Gear {g})'
                series[key].append((float(spd), float(val)))

            if not series:
                continue

            try:
                fig, ax = plt.subplots(figsize=(10, 5))
                fig.patch.set_facecolor('#FAFAFA')
                ax.set_facecolor('#FFFFFF')

                for idx, (label, pts) in enumerate(sorted(series.items())):
                    pts_sorted = sorted(pts, key=lambda x: x[0])
                    xs = [p[0] for p in pts_sorted]
                    ys = [p[1] for p in pts_sorted]
                    color = SLOPE_COLORS[idx % len(SLOPE_COLORS)]
                    ax.plot(xs, ys, color=color, linewidth=1.5, label=label)

                ax.set_xlabel('Speed (km/h)', fontsize=10)
                ax.set_ylabel(ylabel, fontsize=10)
                ax.set_title(chart_title, fontsize=12, fontweight='bold', color='#1E293B')
                ax.grid(True, linestyle='--', alpha=0.5, color='#CBD5E1')
                ax.spines['top'].set_visible(False)
                ax.spines['right'].set_visible(False)
                ax.yaxis.set_major_formatter(mticker.FuncFormatter(
                    lambda x, _: f'{x:,.0f}' if x >= 1000 else f'{x:.1f}'
                ))
                ax.legend(fontsize=7, loc='upper right', framealpha=0.8,
                          ncol=max(1, len(series) // 8))

                buf = io.BytesIO()
                fig.tight_layout()
                plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
                plt.close(fig)
                buf.seek(0)

                doc.add_page_break()
                doc.add_heading(f'7.  {chart_title}', level=1)
                doc.add_picture(buf, width=Inches(6.3))
            except Exception as exc:
                doc.add_paragraph(f'  Chart "{chart_title}" could not be generated: {exc}')

    # ══════════════════════════════════════════════════════════════════
    # SECTION 8 — SUMMARY
    # ══════════════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('8.  Summary', level=1)

    sum_tbl = doc.add_table(rows=1, cols=2)
    sum_tbl.style = 'Table Grid'
    _hdr_row(sum_tbl, 'Result', 'Value')

    summary_rows = [
        ('Locomotive GVW',              f'{gvw_kg:,.0f} kg  ({gvw_ton:.2f} t)'),
        ('Peak Power',                  f'{peak_kw:.0f} kW  ({peak_kw*1.341:.0f} HP)'),
        ('Gear Ratio(s)',               ', '.join(f'{g:.2f}' for g in gear_ratios)),
        ('Rear Axle Ratio',             f'{rear_axle:.3f}'),
        ('Friction Coefficient μ',      f'{mu:.3f}'),
        ('',                            ''),
        ('Max Traction Generated',      f'{gen_n:,.2f} N'),
        ('No-Slip Limit',               f'{slip_n:,.2f} N'),
        ('Effective Traction',          f'{actual_n:,.2f} N'),
        ('Traction Status',             status),
        ('',                            ''),
    ]

    if speed_slope:
        # Max speed on flat
        flat_row = next((r for r in speed_slope if r.get('slope', 1) == 0.0), None)
        if flat_row:
            summary_rows.append(('Max Speed (flat, 0%)',  f"{flat_row['max_speed_kmh']:.2f} km/h"))
        # Speed at max slope
        last_row = speed_slope[-1]
        summary_rows.append((f"Max Speed at {last_row.get('slope',0):.1f}% slope",
                              f"{last_row['max_speed_kmh']:.2f} km/h"))

    for i, (label, value) in enumerate(summary_rows):
        rc = _add_row(sum_tbl, label, value)
        if not label:
            _shade_row(rc, 'F1F5F9')
        elif 'Status' in label:
            fill = 'FEE2E2' if 'slip' in value.lower() else 'DCFCE7'
            _shade_row([rc[1]], fill)
        elif i % 2 == 0:
            _shade_row(rc, 'F8FAFC')

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.add_run(f'— Report generated by Premnathrail Vehicle Performance Calculator  |  '
                 f'{datetime.now().strftime("%d %b %Y")} —').font.size = Pt(8)

    # ── Save ──────────────────────────────────────────────────────────
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream
