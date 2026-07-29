# Vehicle Performance Tool Service Layer
# Orchestrates calculations and prepares formatted reports

import math
import io
from datetime import datetime
from typing import Dict, List, Any

# --- Try to import required libraries ---
try:
    import numpy as np
except ImportError:
    print("="*50)
    print("ERROR: 'numpy' library not found for VehiclePerformance.")
    print("Install it using: pip install numpy")
    print("="*50)
    np = None

try:
    import docx
    # from docx.shared import Inches, Pt
    # from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("="*50)
    print("ERROR: 'python-docx' library not found.")
    print("Install it using: pip install python-docx")
    print("="*50)
    docx = None

# ===============================================================


# Core Physics Functions
# (Extracted from vehicle_performance_calculator.py)
# Units:
# - Speed: km/h
# - Weight: metric tons
# - Forces: Newtons (N)
# - Torque: Newton-meters (Nm), RPM in rev/min, Power in kW
# - Slope: percent (%), curve: degree
# ===============================================================

def rolling_resistance_loco(speed_kmh: float, weight_ton: float, num_axles: int) -> float:
    """Rolling resistance for locomotive in Newtons."""
    if weight_ton <= 0 or num_axles <= 0:
        return 0.0
    A = 0.647 + (13.17 / (weight_ton / num_axles))
    B = 0.00933
    C = 0.057 / weight_ton
    return (A + B * speed_kmh + C * speed_kmh ** 2) * weight_ton * 9.81

def rolling_resistance_wagon(speed_kmh: float, weight_ton: float) -> float:
    """Rolling resistance for wagon in Newtons."""
    if weight_ton <= 0:
        return 0.0
    A = 0.6438797
    B = 0.01047218
    C = 0.00007323
    return (A + B * speed_kmh + C * speed_kmh ** 2) * weight_ton * 9.81

def gradient_resistance(weight_ton: float, slope_percent: float) -> float:
    """Gradient (grade) resistance in Newtons."""
    return weight_ton * 1000 * 9.81 * slope_percent / 100.0

def curvature_resistance(weight_ton: float, curve_degree: float) -> float:
    """Curvature resistance in Newtons."""
    return 0.4 * weight_ton * curve_degree * 9.81

def starting_resistance_loco(weight_ton: float) -> float:
    """Starting resistance for locomotive in Newtons."""
    return 6.0 * weight_ton * 9.81

def starting_resistance_wagon(weight_ton: float) -> float:
    """Starting resistance for wagon in Newtons."""
    return 4.0 * weight_ton * 9.81

def interpolate_torque(engine_rpm: float, curve: Dict[int, float]) -> float:
    """Interpolate engine torque (Nm) at a given RPM from a discrete curve."""
    if not curve:
        raise ValueError("Torque curve is empty.")
    if np is None:
        raise ImportError("Numpy is required for interpolation.")
        
    # Ensure keys are sorted for interpolation
    rpms = np.array(sorted(curve.keys()), dtype=float)
    torques = np.array([curve[r] for r in rpms], dtype=float)
    
    # np.interp handles extrapolation (clamping to endpoints)
    return float(np.interp(engine_rpm, rpms, torques))


# ===============================================================


class VehiclePerformanceCalculator:
    """
    Handles all business logic for Vehicle Performance calculations.
    This class is separate from the FastAPI web server logic.
    """
    
    def __init__(self, inputs: Dict[str, Any]):
        """
        Initializes the calculator with validated and processed inputs
        from the main.py validation function.
        """
        self.inputs = inputs
        
        # --- Store key parameters from the validated inputs dict ---
        # Track
        self.max_curve_deg = inputs.get('max_curve', 0.0)
        if inputs.get('curve_unit') == 'm' and self.max_curve_deg != 0:
            self.max_curve_deg = 1750.0 / self.max_curve_deg
            
        self.max_slope_percent = inputs.get('max_slope', 0.0)
        if inputs.get('slope_unit') == 'degree':
            self.max_slope_percent = math.tan(math.radians(self.max_slope_percent)) * 100.0

        # Vehicle
        self.loco_gvw_kg = inputs.get('loco_gvw_kg', 0.0)
        self.loco_gvw_ton = self.loco_gvw_kg / 1000.0
        self.max_speed_kmh = inputs.get('max_speed_kmh', 0.0)
        self.num_axles = inputs.get('num_axles', 1)
        self.rear_axle_ratio = inputs.get('rear_axle_ratio', 1.0)
        self.gear_ratios = inputs.get('gear_ratios', [1.0])
        self.shunting_load_t = inputs.get('shunting_load_t', 0.0)
        
        # RRV
        self.peak_power_kw = inputs.get('peak_power_kw', 0.0)
        self.friction_mu = inputs.get('friction_mu', 0.3)
        self.wheel_dia_m = inputs.get('wheel_dia_m', 0.73)
        self.min_rpm = inputs.get('min_rpm', 100)
        self.max_rpm = inputs.get('max_rpm', 2500)
        
        # Torque Curve (comes in as Dict[int, float])
        self.torque_curve = inputs.get('torque_curve', {})
        if not self.torque_curve:
            raise ValueError("Torque Curve is empty or missing.")
            
        self.max_torque_nm = max(self.torque_curve.values()) if self.torque_curve else 0.0

    def run_tractive_calculation(self) -> Dict[str, Any]:
        """
        Computes traction vs slipping limits.
        Adapted from Tkinter 'run_tractive_calculation'.
        """
        if self.wheel_dia_m == 0:
             raise ValueError("Wheel Diameter cannot be zero.")
             
        max_traction_generated_n = 2 * (self.max_torque_nm * max(self.gear_ratios) * self.rear_axle_ratio) / self.wheel_dia_m
        max_traction_slipping_n = self.loco_gvw_ton * self.friction_mu * 1000.0 * 9.81
        
        if max_traction_generated_n > max_traction_slipping_n:
            result_message = "Limited by slipping"
        else:
            result_message = "Not limited by slipping"
            
        return {
            "max_traction_generated_n": max(max_traction_generated_n, 0),
            "max_traction_slipping_n": round(max_traction_slipping_n, 2),
            "result_message": result_message
        }

    def _calculate_curves_common(self, calculate_shunting: bool) -> List[Dict[str, Any]]:
        """
        Core plotting engine.
        Adapted from Tkinter 'calculate_common'.
        Returns a list of data points for plotting.
        """
        if np is None:
            raise ImportError("Numpy is required to calculate plot data.")
        
        plot_data: List[Dict[str, Any]] = []
        
        # Iterate over a range of slopes
        for slope_percent in np.arange(0, self.max_slope_percent + 0.5, 0.5):
            for gear_ratio in self.gear_ratios:
                if self.rear_axle_ratio <= 0 or self.wheel_dia_m <= 0 or gear_ratio <= 0:
                    continue

                min_speed_kmh: float = 0.0
                max_speed_kmh: float = float((self.max_rpm * math.pi * self.wheel_dia_m) / (gear_ratio * self.rear_axle_ratio * 60.0) * 3.6)

                # Create a list of speeds (float) without numpy
                if max_speed_kmh <= min_speed_kmh:
                    speeds_kmh = [float(min_speed_kmh)]
                else:
                    step = (max_speed_kmh - min_speed_kmh) / 99.0
                    speeds_kmh = [min_speed_kmh + step * i for i in range(100)]

                for speed_kmh in speeds_kmh:
                    speed_kmh = float(speed_kmh)
                    speed_mps: float = float(speed_kmh / 3.6)
                    wheel_circumference_m: float = float(math.pi * self.wheel_dia_m)
                    wheel_rpm: float = float((speed_mps / wheel_circumference_m * 60.0) if wheel_circumference_m > 0 else 0)
                    engine_rpm: float = float(wheel_rpm * gear_ratio * self.rear_axle_ratio)

                    torque_at_rpm: float = float(interpolate_torque(engine_rpm, self.torque_curve))
                    power_at_rpm_kw: float = float((engine_rpm * torque_at_rpm * 2.0 * math.pi) / (60.0 * 1000.0))

                    if power_at_rpm_kw > self.peak_power_kw and engine_rpm > 0:
                        torque_at_rpm = float((self.peak_power_kw * 60.0 * 1000.0) / (engine_rpm * 2.0 * math.pi))

                    max_traction_generated_n: float = float(2.0 * (torque_at_rpm * gear_ratio * self.rear_axle_ratio) / self.wheel_dia_m)
                    max_traction_slipping_n: float = float(self.loco_gvw_ton * self.friction_mu * 1000.0 * 9.81)
                    actual_traction_n: float = float(min(max_traction_generated_n, max_traction_slipping_n))

                    loco_resistance: float = (
                        float(rolling_resistance_loco(speed_kmh, self.loco_gvw_ton, self.num_axles))
                        + float(gradient_resistance(self.loco_gvw_ton, slope_percent))
                        + float(curvature_resistance(self.loco_gvw_ton, self.max_curve_deg))
                        + float(starting_resistance_loco(self.loco_gvw_ton))
                    )

                    y_val: float = 0.0
                    if calculate_shunting:
                        remaining_traction_n: float = float(actual_traction_n - loco_resistance)
                        if remaining_traction_n > 0:
                            total_resistance_1ton_wagon: float = (
                                float(rolling_resistance_wagon(speed_kmh, 1))
                                + float(gradient_resistance(1, slope_percent))
                                + float(curvature_resistance(1, self.max_curve_deg))
                                + float(starting_resistance_wagon(1))
                            )
                            if total_resistance_1ton_wagon > 0:
                                y_val = float(remaining_traction_n / total_resistance_1ton_wagon)
                    else:
                        y_val = float(actual_traction_n)

                    plot_data.append({
                        "slope": round(float(slope_percent), 2),
                        "gear": gear_ratio,
                        "speed_kmh": round(float(speed_kmh), 2),
                        "value": round(float(y_val), 2)
                    })
        return plot_data

    def calculate_plot_data(self) -> Dict[str, Any]:
        """
        Generates data for both plots expected by the frontend.
        """
        tractive_effort_data = self._calculate_curves_common(calculate_shunting=False)
        shunting_capability_data = self._calculate_curves_common(calculate_shunting=True)
        
        return {
            "tractive_effort_plot": tractive_effort_data,
            "shunting_capability_plot": shunting_capability_data
        }

    def calculate_speed_for_shunting_load(self) -> List[Dict[str, Any]]:
        """
        Compute max achievable speed (km/h) vs slope (%) for a given shunting load.
        Adapted from Tkinter 'calculate_speed_for_shunting_load'.
        """
        if np is None:
            raise ImportError("Numpy is required to calculate plot data.")

        table_data: List[Dict[str, Any]] = []
        
        # Using the tallest ratio for min speed and shortest for max speed
        if self.rear_axle_ratio <= 0 or self.wheel_dia_m <= 0:
             raise ValueError("Axle Ratio or Wheel Diameter cannot be zero.")

        min_speed_kmh = (self.min_rpm * math.pi * self.wheel_dia_m) / (max(self.gear_ratios) * self.rear_axle_ratio * 60.0) * 3.6
        max_speed_kmh = (self.max_rpm * math.pi * self.wheel_dia_m) / (min(self.gear_ratios) * self.rear_axle_ratio * 60.0) * 3.6
        
        if max_speed_kmh <= min_speed_kmh:
            speeds_kmh = [float(min_speed_kmh)]
        else:
            step = (max_speed_kmh - min_speed_kmh) / 99.0
            speeds_kmh = [min_speed_kmh + step * i for i in range(100)]

        for slope_percent in np.arange(0, self.max_slope_percent + 0.5, 0.5):
            max_achievable_speed: float = 0.0

            for speed_kmh in speeds_kmh:
                speed_kmh = float(speed_kmh)
                speed_mps: float = float(speed_kmh / 3.6)
                wheel_circ: float = float(math.pi * self.wheel_dia_m)
                wheel_rpm: float = float((speed_mps / wheel_circ * 60.0) if wheel_circ > 0 else 0)
                engine_rpm: float = float(wheel_rpm * max(self.gear_ratios) * self.rear_axle_ratio)

                tq: float = float(interpolate_torque(engine_rpm, self.torque_curve))

                power_kw: float = float((engine_rpm * tq * 2.0 * math.pi) / (60.0 * 1000.0))
                if power_kw > self.peak_power_kw and engine_rpm > 0:
                    tq = float((self.peak_power_kw * 60.0 * 1000.0) / (engine_rpm * 2.0 * math.pi))

                max_traction_generated_n: float = float(2 * (tq * max(self.gear_ratios) * self.rear_axle_ratio) / self.wheel_dia_m)
                max_traction_slipping_n: float = float(self.loco_gvw_ton * self.friction_mu * 1000.0 * 9.81)
                actual_traction_n: float = float(min(max_traction_generated_n, max_traction_slipping_n))

                loco_res: float = (
                    float(rolling_resistance_loco(speed_kmh, self.loco_gvw_ton, self.num_axles))
                    + float(gradient_resistance(self.loco_gvw_ton, slope_percent))
                    + float(curvature_resistance(self.loco_gvw_ton, self.max_curve_deg))
                )

                wagon_res: float = (
                    float(rolling_resistance_wagon(speed_kmh, self.shunting_load_t))
                    + float(gradient_resistance(self.shunting_load_t, slope_percent))
                    + float(curvature_resistance(self.shunting_load_t, self.max_curve_deg))
                )

                total_resistance_n: float = float(loco_res + wagon_res)

                if actual_traction_n >= total_resistance_n:
                    max_achievable_speed = float(speed_kmh)

            table_data.append({
                "slope": round(float(slope_percent), 2),
                "max_speed_kmh": round(float(max_achievable_speed), 2)
            })
        return table_data
        
    def create_report_docx(self):
        """
        Creates a .docx report and returns it as an in-memory stream.
        Adapted from Tkinter 'export_to_docx'.
        """
        if docx is None:
            raise ImportError("python-docx library is required to generate .docx files.")
            
        doc = docx.Document()
        doc.add_heading("Vehicle Performance Calculator Report", level=1)
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_heading("Input Parameters", level=2)
        
        # Use the original validated inputs for the report
        inputs = self.inputs 
        
        # Inputs table
        inputs_table_data: Dict[str, Any] = {
            "Loco GVW (kg)": inputs.get('loco_gvw_kg'),
            "Max Speed (km/h)": inputs.get('max_speed_kmh'),
            "Number of Axles": inputs.get('num_axles'),
            "Rear Axle Ratio": inputs.get('rear_axle_ratio'),
            "Gear Ratios": ", ".join(map(str, inputs.get('gear_ratios', []))),
            "Shunting Load (tons)": inputs.get('shunting_load_t'),
            "Peak Power (kW)": inputs.get('peak_power_kw'),
            "Coeff. of Friction (mu)": inputs.get('friction_mu'),
            "Wheel Dia (m)": inputs.get('wheel_dia_m'),
            "Min RPM": inputs.get('min_rpm'),
            "Max RPM": inputs.get('max_rpm'),
            "Max Curve": f"{inputs.get('max_curve')} ({inputs.get('curve_unit')})",
            "Max Slope": f"{inputs.get('max_slope')} ({inputs.get('slope_unit')})",
        }
        
        t_inputs = doc.add_table(rows=1, cols=2)
        t_inputs.style = 'Table Grid'
        t_inputs.rows[0].cells[0].text = "Parameter"
        t_inputs.rows[0].cells[1].text = "Value"
        for label, value in inputs_table_data.items():
            row_cells = t_inputs.add_row().cells
            row_cells[0].text = str(label)
            row_cells[1].text = str(value)
            
        # Torque Curve Table
        doc.add_heading("Torque Curve", level=3)
        t_torque = doc.add_table(rows=1, cols=2)
        t_torque.style = 'Table Grid'
        t_torque.rows[0].cells[0].text = "RPM"
        t_torque.rows[0].cells[1].text = "Torque (Nm)"
        for rpm, torque in sorted(self.torque_curve.items()):
            row_cells = t_torque.add_row().cells
            row_cells[0].text = str(rpm)
            row_cells[1].text = str(torque)

        doc.add_heading("Calculation Results", level=2)
        
        # Traction Snapshot
        traction_results = self.run_tractive_calculation()
        doc.add_paragraph(f"Max Traction Produced: {traction_results['max_traction_generated_n']:.2f} N")
        doc.add_paragraph(f"Max Traction (No Slip): {traction_results['max_traction_slipping_n']:.2f} N")
        p = doc.add_paragraph()
        run = p.add_run(f"Result: {traction_results['result_message']}")
        run.font.bold = True

        # Speed vs. Slope Table
        doc.add_heading("Max Achievable Speed vs. Slope", level=3)
        doc.add_paragraph(f"(For Shunting Load: {self.shunting_load_t} tons)")
        
        table_data = self.calculate_speed_for_shunting_load()
        t_speed = doc.add_table(rows=1, cols=2)
        t_speed.style = 'Table Grid'
        t_speed.rows[0].cells[0].text = "Slope (%)"
        t_speed.rows[0].cells[1].text = "Max Achievable Speed (km/h)"
        for row in table_data:
            row_cells = t_speed.add_row().cells
            row_cells[0].text = f"{row['slope']:.2f}"
            row_cells[1].text = f"{row['max_speed_kmh']:.2f}"

        # Save to in-memory stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream


def perform_vehicle_performance_calculation(inputs: Dict[str, Any]) -> Dict[str, Any]:
    """Main calculation function for vehicle performance"""
    calculator = VehiclePerformanceCalculator(inputs)

    results: Dict[str, Any] = {
        'traction_snapshot': calculator.run_tractive_calculation(),
        'plot_data': calculator.calculate_plot_data(),
        'speed_slope_table': calculator.calculate_speed_for_shunting_load()
    }

    return results