import io
import math
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

# Shared templates folder: backend/app/utils/templates/
_TEMPLATES_DIR = Path(__file__).resolve().parents[5] / "utils" / "templates"

_WAGON_RR_START = 4.0
_LOCO_RR_START = 6.0
_WAGON_RR_RUNNING = 1.3505
_LOCO_RR_RUNNING = 2.913
_POWER_CONSTANT = 270
_OHE_VOLTAGE = 22500
_OHE_EFFICIENCY = 0.84
_POWER_FACTOR = 0.8
_CURRENT_CONSTANT = 735.5
_GRADIENT_CONSTANT = 1000
_CURVATURE_CONSTANT = 700


def create_te_docx_report(
    inputs: Dict[str, Any],
    results: Dict[str, Any],
    inputs_raw: Optional[Dict[str, Any]] = None,
) -> io.BytesIO:
    from docx import Document

    if inputs_raw is None:
        inputs_raw = inputs

    doc = Document()
    doc.add_heading("Tractive Effort Calculation Report", 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 1. INPUTS
    doc.add_heading('1. Inputs', level=1)
    load_val = float(inputs.get('load', 0))
    loco_wt_val = float(inputs.get('loco_weight', 0))
    total_wt_val = load_val + loco_wt_val
    gradient_val = float(inputs.get('gradient', 0))
    grad_type = inputs.get('grad_type', 'Degree')
    curvature_val = float(inputs.get('curvature', 0))
    curvature_unit = inputs.get('curvature_unit', 'Radius(m)')
    speed_val = float(inputs.get('speed', 0))
    mode_val = inputs.get('mode', 'Start')

    p = doc.add_paragraph()
    p.add_run(f"  Shunting Load (Load):         {load_val} t\n")
    p.add_run(f"  GBW of Vehicle (Loco Weight): {loco_wt_val} t\n")
    p.add_run(f"  Total Weight:                 {total_wt_val} t\n")
    p.add_run(f"  Gradient:                     {inputs_raw.get('gradient', gradient_val)} ({grad_type})\n")
    p.add_run(f"  Curvature:                    {inputs_raw.get('curvature', curvature_val)} ({curvature_unit})\n")
    p.add_run(f"  Speed:                        {speed_val} km/h\n")
    p.add_run(f"  Mode:                         {mode_val}\n")

    # 2. STEP-BY-STEP CALCULATIONS
    doc.add_heading('2. Step-by-Step Calculations', level=1)

    # Step 1: Gradient Resistance
    doc.add_heading('Step 1: Gradient Resistance (T3)', level=2)
    p = doc.add_paragraph()
    grad_rr_per_ton = results.get('grad_rr_per_ton', 0.0)
    if grad_type == "Degree":
        p.add_run("Formula: R_grad = tan(θ°) × 1000  [kg/t]\n")
        p.add_run(f"  θ = {gradient_val}°\n")
        p.add_run(f"  tan({gradient_val}°) = {math.tan(math.radians(gradient_val)):.6f}\n")
        p.add_run(f"  R_grad = {math.tan(math.radians(gradient_val)):.6f} × {_GRADIENT_CONSTANT} = {grad_rr_per_ton:.4f} kg/t\n")
    else:
        p.add_run("Formula: R_grad = 1000 / G  [kg/t]  (for 1:G gradient)\n")
        p.add_run(f"  G = {gradient_val}\n")
        p.add_run(f"  R_grad = {_GRADIENT_CONSTANT} / {gradient_val} = {grad_rr_per_ton:.4f} kg/t\n")
    T3_val = results.get('T3', total_wt_val * grad_rr_per_ton)
    p.add_run(f"  T3 = {total_wt_val} × {grad_rr_per_ton:.4f} = {T3_val:.2f} kg\n")

    # Step 2: Curvature Resistance
    doc.add_heading('Step 2: Curvature Resistance (T4)', level=2)
    p = doc.add_paragraph()
    curve_rr_per_ton = results.get('curve_rr_per_ton', 0.0)
    if curvature_unit == "Radius(m)":
        p.add_run("Formula: R_curve = 700 / R  [kg/t]\n")
        p.add_run(f"  R = {curvature_val} m\n")
        p.add_run(f"  R_curve = {_CURVATURE_CONSTANT} / {curvature_val} = {curve_rr_per_ton:.4f} kg/t\n")
    else:
        p.add_run("Formula: R_curve = curvature value (direct degree input)\n")
        p.add_run(f"  R_curve = {curve_rr_per_ton:.4f} kg/t\n")
    T4_val = results.get('T4', total_wt_val * curve_rr_per_ton)
    p.add_run(f"  T4 = {total_wt_val} × {curve_rr_per_ton:.4f} = {T4_val:.2f} kg\n")

    # Step 3: Rolling Resistance Coefficients
    doc.add_heading('Step 3: Rolling Resistance Coefficients', level=2)
    p = doc.add_paragraph()
    if mode_val == "Start":
        wagon_rr, loco_rr = _WAGON_RR_START, _LOCO_RR_START
        speed_for_power = 1.0
        p.add_run("  Mode = Start → using starting-mode coefficients\n")
    else:
        wagon_rr, loco_rr = _WAGON_RR_RUNNING, _LOCO_RR_RUNNING
        speed_for_power = speed_val
        p.add_run("  Mode = Running → using running-mode coefficients\n")
    p.add_run(f"  Wagon Rolling Resistance = {wagon_rr} kg/t\n")
    p.add_run(f"  Loco  Rolling Resistance = {loco_rr} kg/t\n")
    p.add_run(f"  Speed for power calc     = {speed_for_power} km/h\n")

    # Step 4: T1–T4
    doc.add_heading('Step 4: Resistance Components T1–T4', level=2)
    T1_val = results.get('T1', load_val * wagon_rr)
    T2_val = results.get('T2', loco_wt_val * loco_rr)
    p = doc.add_paragraph()
    p.add_run("  T1 = Load × Wagon RR   →  Wagon rolling resistance\n")
    p.add_run("  T2 = Loco Wt × Loco RR →  Loco  rolling resistance\n")
    p.add_run("  T3 = Total Wt × R_grad  →  Gradient resistance\n")
    p.add_run("  T4 = Total Wt × R_curve →  Curvature resistance\n\n")
    p.add_run(f"  T1 = {load_val} × {wagon_rr} = {T1_val:.2f} kg\n")
    p.add_run(f"  T2 = {loco_wt_val} × {loco_rr} = {T2_val:.2f} kg\n")
    p.add_run(f"  T3 = {total_wt_val} × {grad_rr_per_ton:.4f} = {T3_val:.2f} kg\n")
    p.add_run(f"  T4 = {total_wt_val} × {curve_rr_per_ton:.4f} = {T4_val:.2f} kg\n")

    # Step 5: Tractive Effort
    doc.add_heading('Step 5: Tractive Effort (TE)', level=2)
    te_val = results.get('te', T1_val + T2_val + T3_val + T4_val)
    p = doc.add_paragraph()
    p.add_run("Formula: TE = T1 + T2 + T3 + T4\n")
    p.add_run(f"  TE = {T1_val:.2f} + {T2_val:.2f} + {T3_val:.2f} + {T4_val:.2f}\n")
    p.add_run(f"  TE = {te_val:.2f} kg  ({te_val/1000:.3f} t)\n").bold = True

    # Step 6: Rail Horsepower
    doc.add_heading('Step 6: Rail Horsepower', level=2)
    power_val = results.get('power', (te_val * speed_for_power) / _POWER_CONSTANT)
    p = doc.add_paragraph()
    p.add_run(f"Formula: Power (HP) = (TE × Speed) / {_POWER_CONSTANT}\n")
    p.add_run(f"  Speed used = {speed_for_power} km/h\n")
    p.add_run(f"  Power = ({te_val:.2f} × {speed_for_power}) / {_POWER_CONSTANT}\n")
    p.add_run(f"  Power = {power_val:.2f} HP\n").bold = True

    # Step 7: OHE Current
    doc.add_heading('Step 7: OHE Current', level=2)
    ohe_current_val = results.get('ohe_current', (power_val * _CURRENT_CONSTANT) / (_OHE_VOLTAGE * _OHE_EFFICIENCY * _POWER_FACTOR))
    p = doc.add_paragraph()
    p.add_run(f"Formula: I = (Power × {_CURRENT_CONSTANT}) / (V_OHE × η_OHE × PF)\n")
    p.add_run(f"  V_OHE = {_OHE_VOLTAGE} V,  η_OHE = {_OHE_EFFICIENCY},  PF = {_POWER_FACTOR}\n")
    p.add_run(f"  I = ({power_val:.2f} × {_CURRENT_CONSTANT}) / ({_OHE_VOLTAGE} × {_OHE_EFFICIENCY} × {_POWER_FACTOR})\n")
    p.add_run(f"  I = {ohe_current_val:.2f} A\n").bold = True

    # 3. RESULTS SUMMARY
    doc.add_heading('3. Results Summary', level=1)
    p = doc.add_paragraph()
    p.add_run("Summary:\n").bold = True
    p.add_run(f"  Tractive Effort (TE) : {te_val:.2f} kg  ({te_val/1000:.3f} t)\n")
    p.add_run(f"  Rail Horsepower      : {power_val:.2f} HP\n")
    p.add_run(f"  OHE Current          : {ohe_current_val:.2f} A\n")
    p = doc.add_paragraph()
    p.add_run("Resistance Components:\n").bold = True
    p.add_run(f"  T1 (Wagon Rolling)   : {T1_val:.2f} kg\n")
    p.add_run(f"  T2 (Loco  Rolling)   : {T2_val:.2f} kg\n")
    p.add_run(f"  T3 (Gradient)        : {T3_val:.2f} kg\n")
    p.add_run(f"  T4 (Curvature)       : {T4_val:.2f} kg\n")
    p.add_run(f"  Total TE             : {te_val:.2f} kg\n")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
