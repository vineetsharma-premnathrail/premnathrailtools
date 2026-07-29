import io
from datetime import datetime
from typing import Dict, Any

try:
    import docx
    from docx.shared import RGBColor
except ImportError:
    docx = None
    RGBColor = None

from ..constants import CONSTANT_C


def create_qmax_docx_report(results: Dict[str, Any], inputs_raw: Dict[str, Any]) -> io.BytesIO:
    if docx is None:
        raise ImportError("python-docx library is required to generate .docx files.")

    doc = docx.Document()
    doc.add_heading("Qmax Calculation Report", 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    if inputs_raw.get('doc_no'):
        doc.add_paragraph(f"Document No: {inputs_raw['doc_no']}")
    if inputs_raw.get('made_by'):
        doc.add_paragraph(
            f"Prepared by: {inputs_raw['made_by']}  |  "
            f"Checked: {inputs_raw.get('checked_by', '')}  |  "
            f"Approved: {inputs_raw.get('approved_by', '')}"
        )

    doc.add_heading("1. Input Parameters", level=1)
    p = doc.add_paragraph()
    p.add_run(f"  • Worn rail diameter limit (d): {inputs_raw.get('d')} mm\n")
    p.add_run(f"  • Material Strength (σB): {inputs_raw.get('sigma_b_selection')}\n")
    p.add_run(f"    (Value Used: {results['sigma_b']} N/mm²)\n")
    p.add_run(f"  • Safety Factor (v_head): {inputs_raw.get('v_head')}\n")

    doc.add_heading("2. Calculation", level=1)
    sigma_v_head_squared = (results['sigma_b'] / results['v_head']) ** 2
    d_half = results['d'] / 2

    doc.add_paragraph(
        f"Formula:   Qmax = C × (d / 2) × (σB / v_head)²\n"
        f"           Where C = {CONSTANT_C}\n\n"
        f"a) (σB / v_head)² = ({results['sigma_b']} / {results['v_head']})² = {sigma_v_head_squared:.3f}\n"
        f"b) Qmax = {CONSTANT_C} × ({results['d']} / 2) × {sigma_v_head_squared:.3f}\n"
        f"c) Qmax = {CONSTANT_C} × {d_half:.1f} × {sigma_v_head_squared:.3f}"
    )

    doc.add_heading("3. Final Result", level=1)
    p_res = doc.add_paragraph()
    run_kn = p_res.add_run(f"Qmax = {results['qmax_kn']:.4f} kN\n")
    run_kn.font.bold = True
    if RGBColor is not None:
        try:
            run_kn.font.color.rgb = RGBColor.from_string("0D47A1")
        except Exception:
            pass
    p_res.add_run(f"Qmax = {results['qmax_tonnes']:.4f} tonnes")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
