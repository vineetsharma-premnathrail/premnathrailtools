import io
from datetime import datetime
from pathlib import Path
from typing import Dict, Any

# Shared templates folder: backend/app/utils/templates/
_TEMPLATES_DIR = Path(__file__).resolve().parents[5] / "utils" / "templates"


def create_load_distro_docx(inputs: Dict[str, Any], results: Dict[str, Any]) -> io.BytesIO:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading('Load Distribution Safety Report', 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading('1. Input Parameters', level=1)
    p = doc.add_paragraph()
    p.add_run(f"  • Configuration Type: {inputs['config_type']}\n")
    p.add_run(f"  • Total Load: {inputs['total_load']:.2f} Ton\n")
    p.add_run(f"  • Front Load Percentage: {inputs['front_percent']:.2f}%\n")
    p.add_run(f"  • Q1 Percentage (of Front Load): {inputs['q1_percent']:.2f}%\n")
    p.add_run(f"  • Q3 Percentage (of Rear Load): {inputs['q3_percent']:.2f}%\n")

    doc.add_heading('2. Calculation Results Summary', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(4.0)
    table.columns[1].width = Inches(2.5)
    cell_text = (
        f"Overall Status: {results['status'].upper()}\n\n"
        f"ΔQ/Q Ratio: {results['delta_q_by_q']:.2%}\n"
        f"Allowed Limit: {results['limit']:.0%}\n\n"
        f"Front Load: {results['front_load']:.2f} Ton\n"
        f"Rear Load: {results['rear_load']:.2f} Ton\n\n"
        f"Q1: {results['q_values']['Q1']:.2f} Ton | Q2: {results['q_values']['Q2']:.2f} Ton\n"
        f"Q3: {results['q_values']['Q3']:.2f} Ton | Q4: {results['q_values']['Q4']:.2f} Ton"
    )
    table.cell(0, 0).text = cell_text

    diagram_path = _TEMPLATES_DIR / "Diagram.png"
    p_img = table.cell(0, 1).paragraphs[0]
    if diagram_path.exists():
        try:
            p_img.add_run().add_picture(str(diagram_path), width=Inches(2.5))
        except Exception as e:
            p_img.text = f"Diagram could not be added: {e}"
    else:
        p_img.text = "Diagram.png not found."

    doc.add_heading('3. Detailed Calculation Steps', level=1)
    from ..service import format_load_distro_steps
    doc.add_paragraph(format_load_distro_steps(inputs, results))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
