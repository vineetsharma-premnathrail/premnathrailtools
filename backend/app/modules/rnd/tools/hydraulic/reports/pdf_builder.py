# Hydraulic Tool Report Builder
# Handles DOCX report generation

import io
from datetime import datetime
from typing import Any, Mapping

# Attempt to import python-docx; add type-ignore so Pylance doesn't complain if it's not installed locally.
try:
    import docx
except Exception:  # pragma: no cover
    docx = None

# Provide typing-only docx imports when running static analyzers in type-check mode
from typing import TYPE_CHECKING
if TYPE_CHECKING:
    # These imports are only used for type hints when available in the environment
    try:
        from docx.shared import RGBColor  # noqa: F401
    except Exception:
        pass

# Import RGBColor if available; otherwise provide a minimal stub to satisfy linters/type-checkers
try:
    from docx.shared import RGBColor
except Exception:  # pragma: no cover
    class _RGBColorStub:
        @staticmethod
        def from_string(s: str) -> Any:
            return None
    RGBColor: Any = _RGBColorStub

# WD_COLOR_INDEX used for highlighting text in the docx. Provide a runtime alias for environments
# where python-docx is not installed so linters/type-checkers don't complain.
try:
    from docx.enum.text import WD_COLOR_INDEX as _wd_color_index_runtime
except Exception:  # pragma: no cover
    class _WDColorStub:
        YELLOW = None
    _wd_color_index_runtime = _WDColorStub

def _add_highlighted_result(paragraph: Any, label: str, value: Any, unit: str = "") -> None:
    """Append a labeled highlighted result run to a paragraph."""
    paragraph.add_run(label)
    val_run: Any = paragraph.add_run(f"{value}")
    try:
        val_run.font.bold = True
        # Highlight may be None for stubbed WD_COLOR_INDEX; guard the attribute assignment
        try:
            val_run.font.highlight_color = _wd_color_index_runtime.YELLOW
        except Exception:
            # Some environments or stubs may not support highlight_color
            pass
    except Exception:
        try:
            val_run.font.bold = True
        except Exception:
            pass
    if unit:
        paragraph.add_run(f" {unit}")


def create_hydraulic_docx_report(
    inputs: Mapping[str, Any],
    results: Mapping[str, Any],
    inputs_raw: Mapping[str, Any]
) -> io.BytesIO:
    """Creates a .docx report and returns it as an in-memory stream."""
    if docx is None:
        raise ImportError("python-docx library is required to generate .docx files.")

    # Create the document and annotate as Any; ignore unknown member type for static analysis
    # Create the document; runtime will raise if docx is not installed
    # Use the concrete Document constructor only at runtime via the module; avoid importing the class symbolically
    doc: Any = docx.Document()

    if inputs['calc_mode'] == "calc_cc":
        _create_displacement_docx(doc, inputs, results, inputs_raw)
    elif inputs['calc_mode'] == 'calc_speed':
        _create_speed_docx(doc, inputs, results, inputs_raw)
    elif inputs['calc_mode'] == 'calc_motor_pressure':
        _create_motor_pressure_docx(doc, inputs, results, inputs_raw)
    elif inputs['calc_mode'] == 'calc_gear':
        _create_gear_docx(doc, inputs, results, inputs_raw)
    else:
        _create_speed_docx(doc, inputs, results, inputs_raw)

    # Save to in-memory stream
    file_stream = io.BytesIO()
    # Call save directly on `doc` (annotated as Any above) so type-checkers won't complain
    # Save with a runtime call; if python-docx is present the method will exist
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def _create_displacement_docx(
    doc: Any,
    inputs: Mapping[str, Any],
    results: Mapping[str, Any],
    raw: Mapping[str, Any]
) -> None:
    """Create a detailed DOCX for Pump & Motor (cc) mode matching the requested template."""
    # Header + document metadata
    doc.add_heading("Pump & Motor (cc) Calculation Report", level=0)
    meta = doc.add_paragraph()
    meta.add_run("Doc-No.: ").bold = True
    meta.add_run(str(raw.get('doc_no') or '—'))
    meta.add_run("    Made By: ").bold = True
    meta.add_run(str(raw.get('made_by') or '—'))
    meta.add_run("\nChecked By: ").bold = True
    meta.add_run(str(raw.get('checked_by') or '—'))
    meta.add_run("    Approved By: ").bold = True
    meta.add_run(str(raw.get('approved_by') or '—'))
    meta.add_run("\nDate: ").bold = True
    meta.add_run(str(raw.get('doc_date') or datetime.now().strftime('%Y-%m-%d')))

    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This document presents the theoretical analysis and calculations for determining the required hydraulic motor displacement and pump displacement for a vehicle propulsion system. "
        "Calculations evaluate wheel speed, resistance forces, torque requirements, motor performance, and hydraulic flow requirements to ensure the system meets the specified operating conditions."
    )

    # Theory & formulas
    doc.add_heading("1. Theory and Formulas", level=2)

    doc.add_paragraph("1.1 Vehicle Speed Conversion")
    doc.add_paragraph("   V (m/s) = V (km/h) × 1000 / 3600")

    doc.add_paragraph("1.2 Wheel RPM")
    doc.add_paragraph("   RPM_wheel = (V / Circumference) × 60")

    doc.add_paragraph("1.3 Resistance Forces")
    doc.add_paragraph("   F_total = F_rolling + F_gradient + F_curve + F_starting")

    doc.add_paragraph("1.4 Torque Requirement")
    doc.add_paragraph("   T = F × r")

    doc.add_paragraph("1.5 Hydraulic Motor Displacement")
    doc.add_paragraph("   Displacement (cc/rev) = (T (kg·cm) × 2π) / (P (kg/cm²) × η_m)")

    doc.add_paragraph("1.6 Flow Rate")
    doc.add_paragraph("   Q (LPM) = (Displacement (cc) × RPM × η_v) / 1000")

    doc.add_paragraph("1.7 Pump Displacement")
    doc.add_paragraph("   Pump_cc/rev = (Flow (LPM) × 1000) / (Pump_RPM × η_p)")

    # Application of theory: Inputs
    doc.add_heading("2. Application of Theory to Vehicle", level=2)

    vehicle_data: Mapping[str, Any] = {
        "Vehicle Weight (t)": raw.get('weight'),
        "Number of Axles": raw.get('axles'),
        "Target Speed (km/h)": raw.get('speed'),
        "Slope (% )": raw.get('slope_percent'),
        "Curve (deg)": inputs.get('curve_degree', 0.0),
        "Wheel Diameter (mm)": raw.get('wheel_diameter'),
        "PTO Ratio": raw.get('pto_gear_ratio'),
        "Axle Ratio": raw.get('axle_gear_box_ratio'),
        "Engine Gear Ratio": raw.get('engine_gear_ratio'),
        "Max Vehicle RPM": raw.get('max_vehicle_rpm')
    }

    hydraulic_data: Mapping[str, Any] = {
        "Total Motors": raw.get('num_motors'),
        "Motors per Axle": raw.get('per_axle_motor'),
        "Pressure (input)": f"{raw.get('pressure')} {raw.get('pressure_unit', 'bar')}",
        "Mechanical Efficiency (%)": raw.get('mech_eff_motor'),
        "Motor Vol Eff (%)": raw.get('vol_eff_motor'),
        "Pump Vol Eff (%)": raw.get('vol_eff_pump')
    }

    _populate_input_table(doc, "Vehicle Inputs", vehicle_data)
    _populate_input_table(doc, "Hydraulic Inputs", hydraulic_data)

    # Step-by-step calculations
    doc.add_heading("3. Calculation Procedure", level=2)

    # Step 1 - Vehicle Speed & Wheel RPM
    doc.add_paragraph("Step 1 — Vehicle Speed & Wheel RPM", style=None)
    speed_kph = float(raw.get('speed') or 0.0)
    speed_mps = float(results.get('speed_mps') or 0.0)
    doc.add_paragraph(f"  Speed conversion: V = {speed_kph} × 1000 / 3600 = {speed_mps:.2f} m/s")

    wheel_dia_mm = float(raw.get('wheel_diameter') or 0.0)
    wheel_circ = float(results.get('wheel_circumference') or 0.0)
    wheel_rpm = float(results.get('wheel_rpm') or 0.0)
    doc.add_paragraph(f"  Wheel circumference: C = {wheel_dia_mm} × π / 1000 = {wheel_circ:.2f} m")
    doc.add_paragraph(f"  Wheel RPM: RPM = (V / C) × 60 = {wheel_rpm:.2f} RPM")

    # --- Resistance & Torque (show only when core computed these keys) ---
    # These keys are produced by calculate_displacement_mode and calculate_motor_pressure_mode
    # (see app/tools/hydraulic/core.py).  Speed-mode does NOT compute traction/torque keys.
    if 'total_resistance' in results:
        # Step 2 - Resistance Forces
        doc.add_paragraph("Step 2 — Resistance Forces", style=None)
        rr = results.get('rolling_resistance', 0.0)
        gr = results.get('gradient_resistance', 0.0)
        cr = results.get('curvature_resistance', 0.0)
        sr = results.get('starting_resistance', 0.0)
        total_res = results.get('total_resistance', 0.0)
        doc.add_paragraph(f"  Rolling resistance: {rr:.2f} kN")
        doc.add_paragraph(f"  Gradient resistance: {gr:.2f} kN")
        doc.add_paragraph(f"  Curve resistance: {cr:.2f} kN")
        doc.add_paragraph(f"  Starting resistance: {sr:.2f} kN")
        doc.add_paragraph(f"  Total resistance: {total_res:.2f} kN")

        # Step 3 - Torque Requirement
        doc.add_paragraph("Step 3 — Torque Requirement", style=None)
        wheel_radius = float(results.get('wheel_radius') or 0.0)
        req_total_torque = float(results.get('required_total_torque') or 0.0)
        per_wheel = float(results.get('per_wheel_torque') or 0.0)
        per_axle = float(results.get('per_axle_torque') or 0.0)
        motor_torque = float(results.get('per_gearbox_input_torque') or 0.0)

        doc.add_paragraph(f"  Wheel radius: r = {wheel_dia_mm}/2000 = {wheel_radius:.2f} m")
        doc.add_paragraph(f"  Total wheel torque: T = {total_res:.2f} kN × 1000 × r = {req_total_torque:.2f} Nm")
        doc.add_paragraph(f"  Per wheel: {per_wheel:.2f} Nm")
        doc.add_paragraph(f"  Per axle: {per_axle:.2f} Nm")
        # Clarified torque labels for DOCX
        doc.add_paragraph(f"  Torque required at axle / after gearbox (per motor): {results.get('per_motor_torque'):.2f} Nm")
        doc.add_paragraph(f"  Torque at motor shaft (gearbox input): {motor_torque:.2f} Nm")
    else:
        # Speed-mode: no traction/torque outputs available from core; skip section
        pass

    # Step 5 - Motor Displacement
    doc.add_paragraph("Step 5 — Motor Displacement", style=None)
    torque_kgcm = float(results.get('per_gearbox_input_torque_kg_cm') or 0.0)
    pressure_kgcm2 = float(results.get('pressure_kg_cm2') or 0.0)
    motor_disp = float(results.get('motor_displacement_cc') or 0.0)

    doc.add_paragraph(f"  Motor Torque (kg·cm) = {torque_kgcm:.2f} kg-cm")
    doc.add_paragraph(f"  Pressure (kg/cm²) = {pressure_kgcm2:.2f} kg/cm²")
    doc.add_paragraph(f"  Motor displacement = (T×2π) / (P×η_m) = {motor_disp:.2f} cc/rev")

    # Suggested standard motor displacement (next higher)
    suggested_motor = results.get('suggested_motor_cc')
    if suggested_motor:
        p_sugg = doc.add_paragraph()
        p_sugg.add_run('Suggested standard motor displacement (next higher): ').bold = True
        p_sugg.add_run(f"{suggested_motor:.0f} cc/rev").bold = True

    # Step 6 - Motor Flow Rate
    doc.add_paragraph("Step 6 — Motor Flow Rate", style=None)
    per_motor_flow = float(results.get('per_motor_flow_rate_lpm') or 0.0)
    doc.add_paragraph(f"  Q (per motor) = {motor_disp:.2f} cc × operating RPM / 1000 / (1/vol_eff) = {per_motor_flow:.2f} LPM")

    # Per-motor summary
    doc.add_paragraph(f"  Number of motors: {results.get('num_motors')}")
    doc.add_paragraph(f"  Motors per axle: {results.get('motors_per_axle')}")
    doc.add_paragraph(f"  Torque required at axle / after gearbox (per motor): {results.get('per_motor_torque'):.2f} Nm")

    # Step 7 - Pump Displacement / Table
    doc.add_paragraph("Step 7 — Required Pump Displacement", style=None)
    doc.add_heading("Pump Requirement Table", level=3)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Engine Gear'
    hdr[1].text = 'Max RPM'
    hdr[2].text = 'Pump RPM'
    hdr[3].text = 'Pump Disp (cc/rev)'

    pump_results_list = results.get('pump_results', [])
    if not pump_results_list:
        row = table.add_row().cells
        row[0].text = '-'
        row[1].text = '-'
        row[2].text = '-'
        row[3].text = '-'
    for pr in pump_results_list:
        r = table.add_row().cells
        r[0].text = f"{pr.get('engine_gear_ratio', 0):.2f}"
        r[1].text = f"{pr.get('max_vehicle_rpm_input', 0):.0f}"
        r[2].text = f"{pr.get('pump_rpm', 0):.2f}"
        r[3].text = f"{pr.get('pump_disp_cc', 0):.2f}"

    # Suggested pump displacement (per pump) — show for the first result as example
    if pump_results_list:
        first_pr = pump_results_list[0]
        suggested_pp = first_pr.get('suggested_pump_disp_per_pump_cc') or first_pr.get('suggested_pump_disp_cc')
        if suggested_pp:
            doc.add_paragraph(f"Suggested standard pump displacement per pump (next higher): {suggested_pp:.0f} cc")

    # Summary
    doc.add_heading("4. Summary", level=2)
    # create a table for key output parameters
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'
    hdr = summary_table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Value'

    def _add_summary_row(param: str, val: str) -> None:
        r = summary_table.add_row().cells
        r[0].text = param
        r[1].text = val

    _add_summary_row('Required motor displacement (cc/rev)', f"{motor_disp:.2f}")
    if suggested_motor:
        _add_summary_row('Suggested motor displacement (cc/rev)', f"{suggested_motor:.0f}")
    _add_summary_row('Each motor flow (LPM)', f"{per_motor_flow:.2f}")
    if pump_results_list:
        first_pr = pump_results_list[0]
        _add_summary_row('Pump displacement (example) (cc/rev)', f"{first_pr.get('pump_disp_cc'):.2f}")
        suggested_pp = first_pr.get('suggested_pump_disp_per_pump_cc') or first_pr.get('suggested_pump_disp_cc')
        if suggested_pp:
            _add_summary_row('Suggested pump displacement per pump (cc/rev)', f"{suggested_pp:.0f}")
        total_pump_power = first_pr.get('pump_total_power_kw') or 0
        _add_summary_row('Pump power (total) (kW)', f"{total_pump_power:.2f}")
    # resistances if available
    if results.get('total_resistance') is not None:
        _add_summary_row('Total system resistance (kN)', f"{results.get('total_resistance'):.2f}")

    # include a notes paragraph after table
    doc.add_paragraph("Notes: All formulas and step-by-step calculations are included. Values shown are calculated from inputs and engineering assumptions used in the engineered model.")

def _create_speed_docx(
    doc: Any,
    inputs: Mapping[str, Any],
    results: Mapping[str, Any],
    raw: Mapping[str, Any]
) -> None:
    """DOCX for Speed mode — mirrors the LaTeX PDF structure exactly."""
    doc.add_heading("Hydraulic Speed Calculation Report", level=0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # ── Input Parameters ──────────────────────────────────────────────────────
    doc.add_heading("Input Parameters", level=1)

    doc.add_heading("Vehicle Inputs", level=2)
    vehicle_items = [
        ("Vehicle Weight",       f"{raw.get('weight')} t"),
        ("Total Axles",          raw.get('axles')),
        ("Drive Axles",          raw.get('drive_axles')),
        ("Wheel Diameter",       f"{raw.get('wheel_diameter')} mm"),
        ("Slope",                f"{raw.get('slope_percent')} %"),
        ("Curve",                f"{inputs.get('curve_degree', 0.0)} deg"),
        ("Axle Gear Box Ratio",  raw.get('axle_gear_box_ratio')),
        ("Engine Gear Ratio(s)", raw.get('engine_gear_ratio')),
        ("Max Vehicle RPM",      f"{raw.get('max_vehicle_rpm')} RPM"),
    ]
    for label, value in vehicle_items:
        doc.add_paragraph(f"  {label}: {value}")

    doc.add_heading("Hydraulic Inputs", level=2)
    hydraulic_items = [
        ("PTO Gear Ratio",              raw.get('pto_gear_ratio')),
        ("Number of Pumps",             raw.get('num_pumps')),
        ("Pump Displacement",           f"{raw.get('pump_disp_in')} cc/rev"),
        ("Pump Volumetric Efficiency",  f"{raw.get('vol_eff_pump')} %"),
        ("Motor Displacement",          f"{raw.get('motor_disp_in')} cc/rev"),
        ("Motor Mechanical Efficiency", f"{raw.get('mech_eff_motor')} %"),
        ("System Pressure",             f"{raw.get('pressure')} bar"),
        ("Motors per Axle",             raw.get('per_axle_motor')),
        ("Number of Motors",            raw.get('num_motors')),
    ]
    for label, value in hydraulic_items:
        doc.add_paragraph(f"  {label}: {value}")

    # ── Gear-wise Hydraulic Calculations ─────────────────────────────────────
    doc.add_heading("Gear-wise Hydraulic Calculations", level=1)
    gear_results = results.get('gear_results', [])
    pump_cc   = float(inputs.get('pump_disp_in') or 0)
    motor_cc  = float(inputs.get('motor_disp_in') or 0)
    vol_eff   = float(inputs.get('vol_eff_pump') or 0) / 100.0
    mech_eff  = float(inputs.get('mech_eff_motor') or 0) / 100.0
    pressure  = float(inputs.get('pressure') or 0)
    max_rpm   = float(inputs.get('max_vehicle_rpm') or 0)
    pto       = float(inputs.get('pto_gear_ratio') or 1)
    num_pumps = float(inputs.get('num_pumps') or 1)
    d_axles   = float(inputs.get('drive_axles') or 1)
    mpa       = float(inputs.get('per_axle_motor') or 1)

    for g in gear_results:
        gr = g.get('gear_ratio', 0)
        doc.add_heading(f"Gear Ratio = {gr:.2f}", level=2)

        p = doc.add_paragraph("Step 1: Pump Speed\n")
        p.add_run(f"  Pump_Speed = (Max_RPM / Gear_Ratio) x (PTO_Ratio / Num_Pumps)\n")
        p.add_run(f"  = ({max_rpm:.0f} / {gr:.2f}) x ({pto:.2f} / {num_pumps:.0f})\n")
        _add_highlighted_result(p, "  = ", g.get('pump_speed', 0), "RPM")

        p = doc.add_paragraph("Step 2: Per Pump Flow\n")
        p.add_run(f"  Q_pump = (Pump_cc x Pump_Speed x eta_vol) / 1000\n")
        p.add_run(f"  = ({pump_cc:.1f} x {g.get('pump_speed', 0):.2f} x {vol_eff:.2f}) / 1000\n")
        _add_highlighted_result(p, "  = ", g.get('per_pump_flow', 0), "LPM")

        p = doc.add_paragraph("Step 3: Per Motor Flow\n")
        p.add_run(f"  Q_motor = Q_pump / (Drive_Axles x Motors_per_Axle)\n")
        p.add_run(f"  = {g.get('per_pump_flow', 0):.2f} / ({d_axles:.0f} x {mpa:.0f})\n")
        _add_highlighted_result(p, "  = ", g.get('per_motor_flow', 0), "LPM")

        p = doc.add_paragraph("Step 4: Available Motor Torque\n")
        p.add_run(f"  T_avail = eta_mech x Pressure x Motor_cc x 0.015915\n")
        p.add_run(f"  = {mech_eff:.2f} x {pressure:.1f} x {motor_cc:.1f} x 0.015915\n")
        _add_highlighted_result(p, "  = ", g.get('avail_torque', 0), "Nm")

    # ── Matched Speed Section OR No-Match Section ─────────────────────────────
    matched_spd  = results.get('matched_speed_kph')
    matched_gear = results.get('matched_gear_ratio')

    if matched_spd is not None:
        doc.add_heading("Resistance & Torque at Matched Speed", level=1)
        doc.add_heading(f"Matched Condition: {matched_spd} km/h @ Gear Ratio {matched_gear:.2f}", level=2)

        A = results.get('coeff_A', 0.0)
        B = results.get('coeff_B', 0.0)
        C = results.get('coeff_C', 0.0)
        W = float(inputs.get('weight') or 0)
        axles = float(inputs.get('axles') or 1)

        p = doc.add_paragraph("Step 1: Resistance Coefficients\n")
        p.add_run(f"  A = 0.647 + (13.17 / (W / Axles)) = 0.647 + (13.17 / ({W} / {axles:.0f})) = {A:.2f}\n")
        p.add_run(f"  B = 0.00933\n")
        p.add_run(f"  C = 0.057 / W = 0.057 / {W} = {C:.2f}\n")

        p = doc.add_paragraph(f"Step 2: Resistance Forces at V = {matched_spd} km/h\n")
        p.add_run(f"  Rolling    = {results.get('rolling_resistance', 0):.2f} kN\n")
        p.add_run(f"  Gradient   = {results.get('gradient_resistance', 0):.2f} kN\n")
        p.add_run(f"  Curvature  = {results.get('curvature_resistance', 0):.2f} kN\n")
        p.add_run(f"  Starting   = {results.get('starting_resistance', 0):.2f} kN\n")
        r = p.add_run(f"  F_total    = {results.get('total_resistance', 0):.2f} kN\n")
        r.bold = True

        wr = results.get('wheel_radius', 0)
        p = doc.add_paragraph("Step 3: Torque Requirement\n")
        p.add_run(f"  r = wheel_dia / 2000 = {float(inputs.get('wheel_diameter') or 0):.0f} / 2000 = {wr:.2f} m\n")
        p.add_run(f"  T_total    = {results.get('required_total_torque', 0):.2f} Nm\n")
        p.add_run(f"  T_per_axle = {results.get('per_axle_torque', 0):.2f} Nm\n")
        p.add_run(f"  T_per_motor (req) = {results.get('per_gearbox_input_torque', 0):.2f} Nm\n")

        p = doc.add_paragraph("Step 4: Torque Match Check\n")
        avail = results.get('available_torque', 0)
        req   = results.get('per_gearbox_input_torque', 0)
        p.add_run(f"  T_avail = {avail:.2f} Nm  >=  T_req = {req:.2f} Nm  [MATCH]\n")
        r = p.add_run(f"  Matched Speed = {matched_spd} km/h @ Gear Ratio {matched_gear:.2f}")
        r.bold = True
        try:
            r.font.color.rgb = RGBColor.from_string("0D47A1")
        except Exception:
            pass

    else:
        doc.add_heading("No Achievable Speed Found", level=1)
        avail   = results.get('available_torque', 0)
        min_req = results.get('min_required_torque', 0)
        p = doc.add_paragraph()
        p.add_run(f"  T_avail = {avail:.2f} Nm  <  T_req (min at 1 km/h) = {min_req:.2f} Nm\n\n")
        p.add_run(
            f"  At the given slope ({inputs.get('slope_percent', 0)} %), "
            f"curve ({inputs.get('curve_degree', 0)} deg) and gear ratio(s) "
            f"{inputs.get('engine_gear_ratio_list', '---')}, "
            f"the available motor torque is insufficient to move the vehicle "
            f"at any speed from 1-120 km/h.\n\n"
        )
        p.add_run("  Suggestions to achieve a match:\n")
        p.add_run(f"    - Increase system pressure (currently {inputs.get('pressure', 0)} bar)\n")
        p.add_run(f"    - Increase motor displacement (currently {inputs.get('motor_disp_in', 0)} cc/rev)\n")
        p.add_run( "    - Reduce slope or vehicle weight\n")

def _create_motor_pressure_docx(
    doc: Any,
    inputs: Mapping[str, Any],
    results: Mapping[str, Any],
    raw: Mapping[str, Any]
) -> None:
    """DOCX for Motor Pressure mode — detailed step-by-step with formulas and numeric substitution."""
    doc.add_heading("Motor Pressure Calculation Report", level=0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # ── Input Parameters ──────────────────────────────────────────────────────
    doc.add_heading("Input Parameters", level=1)
    vehicle_data: Mapping[str, Any] = {
        "Vehicle Weight (t)": raw.get('weight'),
        "Drive Axles": raw.get('drive_axles'),
        "Target Speed (km/h)": raw.get('speed'),
        "Wheel Dia (mm)": raw.get('wheel_diameter'),
        "Slope": f"{raw.get('slope_percent')} % (unit: {raw.get('slope_unit', '%')})",
        "Curve (deg)": f"{inputs.get('curve_degree', 0.0):.4f}°",
        "Axle Gear Box Ratio": raw.get('axle_gear_box_ratio'),
    }
    motor_data: Mapping[str, Any] = {
        "Total Motors": raw.get('num_motors'),
        "Motor / Axle": raw.get('per_axle_motor'),
        "Motor Displacement (cc)": raw.get('motor_disp_in'),
        "Motor Mech Eff (%)": raw.get('mech_eff_motor'),
    }
    _populate_input_table(doc, "Vehicle Inputs", vehicle_data)
    _populate_input_table(doc, "Motor Inputs", motor_data)

    # ── Theory ────────────────────────────────────────────────────────────────
    doc.add_heading("Formulas Used", level=1)
    p = doc.add_paragraph()
    p.add_run("Step 1  speed_mps = speed_kph / 3.6\n")
    p.add_run("        circumference = π × D_mm / 1000\n")
    p.add_run("        wheel_rpm = (speed_mps / circumference) × 60\n\n")
    p.add_run("Step 2  Rolling Res  = (A + B×v + C×v²) × W × g / 1000  [kN]\n")
    p.add_run("        Gradient Res = W×1000×g × slope% / 100 000  [kN]\n")
    p.add_run("        Curve Res    = 0.4 × W × curve_deg × g / 1000  [kN]\n")
    p.add_run("        Starting Res = 6 × W × g / 1000  [kN]\n\n")
    p.add_run("Step 3  Total Torque = F_total_kN × 1000 × wheel_radius  [Nm]\n")
    p.add_run("        per_axle_torque  = total_torque / drive_axles\n")
    p.add_run("        per_motor_torque = per_axle_torque / motors_per_axle\n")
    p.add_run("        motor_shaft_torque = per_motor_torque / gear_ratio  [Nm]\n\n")
    p.add_run("Step 4  T_kg_cm = T_Nm × 10.1972  (unit conversion)\n")
    p.add_run("        P_bar = (T_kg_cm × 2π) / (D_cc × η_mech)\n")

    # ── Step-by-Step Calculation ───────────────────────────────────────────────
    doc.add_heading("Step-by-Step Calculation", level=1)

    speed_kph  = float(inputs.get('speed') or 0.0)
    wheel_d_mm = float(inputs.get('wheel_diameter') or 0.0)
    speed_mps  = results.get('speed_mps') or (speed_kph / 3.6)
    circ       = results.get('wheel_circumference') or (3.14159265 * wheel_d_mm / 1000)
    wheel_rpm  = results.get('wheel_rpm') or 0.0

    doc.add_heading("Step 1: Vehicle Speed & Wheel RPM", level=2)
    p = doc.add_paragraph()
    p.add_run(f"  speed_mps = {speed_kph} km/h ÷ 3.6 = {speed_mps:.4f} m/s\n")
    p.add_run(f"  circumference = π × {wheel_d_mm} / 1000 = {circ:.4f} m\n")
    p.add_run(f"  wheel_rpm = ({speed_mps:.4f} / {circ:.4f}) × 60 = ")
    r = p.add_run(f"{wheel_rpm:.2f} RPM\n")
    r.bold = True

    doc.add_heading("Step 2: Resistance Forces", level=2)
    roll  = results.get('rolling_resistance') or 0.0
    grad  = results.get('gradient_resistance') or 0.0
    curve = results.get('curvature_resistance') or 0.0
    start = results.get('starting_resistance') or 0.0
    total = results.get('total_resistance') or 0.0
    p = doc.add_paragraph()
    p.add_run(f"  Rolling Resistance  = {roll:.4f} kN\n")
    p.add_run(f"  Gradient Resistance = {grad:.4f} kN\n")
    p.add_run(f"  Curvature Resistance= {curve:.4f} kN\n")
    p.add_run(f"  Starting Resistance = {start:.4f} kN\n")
    p.add_run("  " + "─" * 42 + "\n")
    r = p.add_run(f"  Total Resistance    = {total:.4f} kN\n")
    r.bold = True

    doc.add_heading("Step 3: Torque Requirements", level=2)
    w_radius      = results.get('wheel_radius') or (wheel_d_mm / 2000)
    tot_torque    = results.get('required_total_torque') or 0.0
    per_whl       = results.get('per_wheel_torque') or 0.0
    per_axle_t    = results.get('per_axle_torque') or 0.0
    per_motor_t   = results.get('per_motor_torque') or 0.0
    gearbox_in_t  = results.get('per_gearbox_input_torque') or 0.0
    p = doc.add_paragraph()
    p.add_run(f"  Wheel Radius        = {wheel_d_mm} / 2000 = {w_radius:.4f} m\n")
    p.add_run(f"  Total Torque        = {total:.4f} kN × 1000 × {w_radius:.4f} = {tot_torque:.2f} Nm\n")
    p.add_run(f"  Per-Wheel Torque    = {tot_torque:.2f} / (2 × drive_axles)  = {per_whl:.2f} Nm\n")
    p.add_run(f"  Per-Axle Torque     = {tot_torque:.2f} / drive_axles = {per_axle_t:.2f} Nm\n")
    p.add_run(f"  Per-Motor Torque    = {per_axle_t:.2f} / motors_per_axle = {per_motor_t:.2f} Nm\n")
    r = p.add_run(f"  Motor Shaft Torque  = {per_motor_t:.2f} / gear_ratio = {gearbox_in_t:.2f} Nm\n")
    r.bold = True

    doc.add_heading("Step 4: Required Pressure", level=2)
    t_kg_cm  = results.get('per_gearbox_input_torque_kg_cm') or 0.0
    d_cc     = float(inputs.get('motor_disp_in') or results.get('motor_displacement_cc') or 0.0)
    eta_mech = float(inputs.get('mech_eff_motor') or 0.0) / 100.0
    req_p    = results.get('required_pressure_bar') or 0.0
    p = doc.add_paragraph()
    p.add_run(f"  T_kg_cm = {gearbox_in_t:.4f} Nm × 10.1972 = {t_kg_cm:.4f} kg·cm\n")
    p.add_run(f"  Formula : P = (T_kg·cm × 2π) / (D_cc × η_mech)\n")
    p.add_run(f"  P = ({t_kg_cm:.4f} × 2π) / ({d_cc:.2f} × {eta_mech:.4f})\n")
    r = p.add_run(f"  Required Pressure = {req_p:.2f} bar\n")
    r.bold = True

    # ── Results Summary ────────────────────────────────────────────────────────
    doc.add_heading("Results Summary", level=1)
    p = doc.add_paragraph()
    p.add_run(f"  Wheel RPM                        = {wheel_rpm:.2f} RPM\n")
    p.add_run(f"  Total Resistance                 = {total:.2f} kN\n")
    p.add_run(f"  Motor Shaft Torque (gearbox in)  = {gearbox_in_t:.2f} Nm\n")
    p2 = doc.add_paragraph("  Required Pressure: ")
    run = p2.add_run(f"{req_p:.2f} bar")
    run.font.bold = True
    try:
        run.font.color.rgb = RGBColor.from_string("0D47A1")
    except Exception:
        pass

    # Suggested standard motor size
    motor_disp_input = float(inputs.get('motor_disp_in') or 0.0)
    if motor_disp_input > 0:
        motor_standard_sizes = [16, 23, 45, 56, 63, 80, 90, 107, 125]
        suggested_motor = next((s for s in motor_standard_sizes if s >= motor_disp_input), motor_standard_sizes[-1])
        doc.add_paragraph(f"  Suggested standard motor (next higher): {suggested_motor:.0f} cc/rev")


def _create_gear_docx(
    doc: Any,
    inputs: Mapping[str, Any],
    results: Mapping[str, Any],
    raw: Mapping[str, Any]
) -> None:
    """DOCX for Gear Ratio mode — detailed step-by-step with formulas and numeric substitution."""
    doc.add_heading("Gear Ratio Calculation Report", level=0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # ── Introduction ──────────────────────────────────────────────────────────
    doc.add_paragraph(
        "This report calculates the required axle gear ratio so that the hydraulic motor "
        "can drive the wheel at the specified target speed without exceeding its maximum RPM."
    )

    # ── Inputs ────────────────────────────────────────────────────────────────
    doc.add_heading("Input Parameters", level=1)
    vehicle_data: Mapping[str, Any] = {
        "Target Speed (km/h)": raw.get('speed'),
        "Wheel Diameter (mm)": raw.get('wheel_diameter'),
    }
    motor_data: Mapping[str, Any] = {
        "Max Motor RPM": raw.get('max_motor_rpm'),
    }
    _populate_input_table(doc, "Vehicle Inputs", vehicle_data)
    _populate_input_table(doc, "Motor Inputs", motor_data)

    # ── Theory ────────────────────────────────────────────────────────────────
    doc.add_heading("Theory & Formulas", level=1)
    p = doc.add_paragraph()
    p.add_run("Step 1  circumference (m) = π × wheel_diameter_mm / 1000\n")
    p.add_run("Step 2  speed_mps = target_speed_kph / 3.6\n")
    p.add_run("        wheel_rpm = (speed_mps / circumference) × 60\n")
    p.add_run("Step 3  required_gear_ratio = max_motor_rpm / wheel_rpm\n")

    # ── Step-by-Step Calculation ───────────────────────────────────────────────
    doc.add_heading("Step-by-Step Calculation", level=1)

    speed_kph    = float(inputs.get('speed') or raw.get('speed') or 0.0)
    wheel_d_mm   = float(inputs.get('wheel_diameter') or raw.get('wheel_diameter') or 0.0)
    max_mot_rpm  = float(inputs.get('max_motor_rpm') or raw.get('max_motor_rpm') or 0.0)
    circ         = results.get('wheel_circumference') or (3.14159265 * wheel_d_mm / 1000)
    wheel_rpm    = results.get('wheel_rpm') or 0.0
    gear_ratio   = results.get('required_gear_ratio') or 0.0
    speed_mps    = speed_kph / 3.6

    doc.add_heading("Step 1: Wheel Circumference", level=2)
    p = doc.add_paragraph()
    p.add_run("  circumference = π × wheel_diameter_mm / 1000\n")
    p.add_run(f"  = 3.14159 × {wheel_d_mm} / 1000\n")
    r = p.add_run(f"  = {circ:.4f} m\n")
    r.bold = True

    doc.add_heading("Step 2: Wheel RPM at Target Speed", level=2)
    p = doc.add_paragraph()
    p.add_run(f"  speed_mps = {speed_kph} km/h ÷ 3.6 = {speed_mps:.4f} m/s\n")
    p.add_run(f"  wheel_rpm = (speed_mps / circumference) × 60\n")
    p.add_run(f"  = ({speed_mps:.4f} / {circ:.4f}) × 60\n")
    r = p.add_run(f"  = {wheel_rpm:.2f} RPM\n")
    r.bold = True

    doc.add_heading("Step 3: Required Gear Ratio", level=2)
    p = doc.add_paragraph()
    p.add_run("  required_gear_ratio = max_motor_rpm / wheel_rpm\n")
    p.add_run(f"  = {max_mot_rpm:.0f} / {wheel_rpm:.2f}\n")
    r = p.add_run(f"  = {gear_ratio:.3f}\n")
    r.bold = True

    # ── Results Summary ────────────────────────────────────────────────────────
    doc.add_heading("Results Summary", level=1)
    p = doc.add_paragraph()
    p.add_run(f"  Wheel Circumference      = {circ:.4f} m\n")
    p.add_run(f"  Wheel RPM at target speed= {wheel_rpm:.2f} RPM\n")
    p2 = doc.add_paragraph("  Required Gear Ratio: ")
    run = p2.add_run(f"{gear_ratio:.3f}")
    run.font.bold = True
    try:
        run.font.color.rgb = RGBColor.from_string("0D47A1")
    except Exception:
        pass
    if results.get('warnings'):
        doc.add_heading("Warnings", level=2)
        for w in results['warnings']:
            doc.add_paragraph(f"  ⚠ {w}")


def _populate_input_table(
    doc: Any,
    title: str,
    data_dict: Mapping[str, Any]
) -> None:
    """Helper to populate input parameter tables"""
    if docx is None:
        return
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=len(data_dict), cols=2)
    table.style = 'Table Grid'
    for i, (key, value) in enumerate(data_dict.items()):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = str(value)
        table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    return table


# --- PDF (LaTeX) generator using Jinja2 + pdflatex ---
def generate_hydraulic_pdf_report(context: Mapping[str, Any]) -> io.BytesIO:
    """Render LaTeX Jinja2 template (services/templates/hydraulic_cc_template.tex) and compile to PDF.

    Improved behavior:
    - Treat produced PDF as success even if the LaTeX process exits with a non-zero code (MiKTeX/pdfTeX sometimes returns 1 but still writes a PDF).
    - Include both stdout/stderr and the .log tail in error messages to make debugging easier.
    - Try pdflatex, then xelatex, then lualatex; accept the first run that produces a PDF file.
    """
    try:
        from jinja2 import Environment, FileSystemLoader
        import os
        from pathlib import Path
        # reports/ → hydraulic/ → tools/ → rnd/ → v1/ → api/ → app/ → utils/templates
        template_dir = str(Path(__file__).resolve().parent.parent.parent.parent.parent.parent / 'utils' / 'templates')
        env = Environment(loader=FileSystemLoader(template_dir))
        mode = (context or {}).get('calc_mode', 'calc_cc')
        if mode == 'calc_speed':
            template = env.get_template('hydraulic_speed_template.tex')
        elif mode == 'calc_motor_pressure':
            template = env.get_template('hydraulic_pressure_template.tex')
        elif mode == 'calc_gear':
            template = env.get_template('hydraulic_gear_template.tex')
        else:
            # default / displacement / other modes use the cc template
            template = env.get_template('hydraulic_cc_template.tex')
        # copy context into a mutable dict and enrich with helpful suggestion keys for templates
        ctx = dict(context or {})
        if mode == 'calc_speed':
            motor_disp_in = float(ctx.get('motor_disp_in') or 0)
            if motor_disp_in:
                motor_standard_sizes = [16, 23, 45, 56, 63, 80, 90, 107, 125]
                suggested_motor = next((s for s in motor_standard_sizes if s >= motor_disp_in), motor_standard_sizes[-1])
                ctx['suggested_motor_cc'] = float(suggested_motor)
            pump_disp_in = float(ctx.get('pump_disp_in') or 0)
            if pump_disp_in:
                pump_standard_sizes = [10, 12, 16, 23, 28, 32, 45, 56, 63, 80, 90, 107, 125, 160, 180]
                suggested_pump = next((s for s in pump_standard_sizes if s >= pump_disp_in), pump_standard_sizes[-1])
                ctx['suggested_pump_disp_per_pump_cc'] = float(suggested_pump)
        if mode == 'calc_motor_pressure':
            motor_disp_in = float(ctx.get('motor_disp_in') or ctx.get('motor_displacement_cc') or 0)
            if motor_disp_in:
                motor_standard_sizes = [16, 23, 45, 56, 63, 80, 90, 107, 125]
                suggested_motor = next((s for s in motor_standard_sizes if s >= motor_disp_in), motor_standard_sizes[-1])
                ctx['suggested_motor_cc'] = float(suggested_motor)
        # calc_cc: motor flow formula in template uses max_motor_rpm as motor operating RPM = wheel_rpm
        if mode == 'calc_cc':
            ctx.setdefault('max_motor_rpm', ctx.get('wheel_rpm', 0))
        # calc_speed: template uses engine_gear_ratio_list for display
        if mode == 'calc_speed' and 'engine_gear_ratio_list' not in ctx:
            ctx['engine_gear_ratio_list'] = ctx.get('engine_gear_ratio', '---')
        latex_content = template.render(**ctx)

        # Compile LaTeX to PDF in temporary dir
        import tempfile
        from pathlib import Path
        import subprocess
        import shutil
        with tempfile.TemporaryDirectory() as td:
            tmp = Path(td)
            tex_file = tmp / 'report.tex'
            tex_file.write_text(latex_content, encoding='utf-8')

            # Copy logos so LaTeX can find them (resolve relative to template_dir)
            tpl_path = Path(template_dir)
            for img in ['logo.JPG', 'logo-1.JPG']:
                src = tpl_path / img
                if src.exists():
                    shutil.copy(src, tmp / img)

            # openin_any=p ("paranoid") restricts kpathsea's \input/\include
            # file resolution to the current directory and its subdirs — set
            # as an env var since kpathsea reads it there in preference to
            # texmf.cnf. Without this, a value like `\input{/app/.env}` in an
            # escaped-but-still-backslash-containing string wouldn't execute
            # (escape_latex neutralizes the backslash), but this is defense
            # in depth against any future field that reaches the template
            # unescaped.
            _tex_env = {**os.environ, "openin_any": "p"}

            def _run(cmd: list[str]) -> subprocess.CompletedProcess[str]:
                return subprocess.run(cmd, cwd=td, capture_output=True, text=True, timeout=60, env=_tex_env)

            # Resolve full paths for engines — checks PATH first, then common Windows install locations
            _win_miktex_dirs = [
                r'C:\Program Files\MiKTeX\miktex\bin\x64',
                r'C:\Program Files\MiKTeX\miktex\bin',
                r'C:\Program Files (x86)\MiKTeX\miktex\bin\x64',
                r'C:\Users\itsvi\AppData\Local\Programs\MiKTeX\miktex\bin\x64',
            ]
            def _resolve_engine(name: str) -> str:
                found = shutil.which(name)
                if found:
                    return found
                for d in _win_miktex_dirs:
                    candidate = os.path.join(d, name + '.exe')
                    if os.path.exists(candidate):
                        return candidate
                return name  # fallback — will raise FileNotFoundError if not found

            engines = [_resolve_engine(e) for e in ['pdflatex', 'xelatex', 'lualatex']]
            last_res: subprocess.CompletedProcess[str] | None = None
            engines_tried: list[str] = []
            engines_missing: list[str] = []
            pdf_path = tmp / 'report.pdf'

            for engine in engines:
                try:
                    last_res = _run([engine, '-interaction=nonstopmode', '-no-shell-escape', 'report.tex'])
                    engines_tried.append(engine)
                except FileNotFoundError:
                    engines_missing.append(engine)
                    last_res = None
                    continue

                # If a PDF was produced, accept it even if returncode != 0 (MiKTeX/pdfTeX behavior)
                if pdf_path.exists():
                    data = pdf_path.read_bytes()
                    return io.BytesIO(data)

                # continue to try next engine
                continue

            # Distinguish "not installed" from "compilation error"
            if len(engines_missing) == len(engines):
                raise Exception(
                    'LaTeX is not installed on this server. '
                    'Install MiKTeX (https://miktex.org/download) on Windows or '
                    'TeX Live on Linux/Mac, then restart the server.'
                )

            # LaTeX was found but compilation failed — include diagnostic info
            diag: list[str] = []
            if last_res is not None:
                diag.append(f"engine={engines_tried[-1]}  returncode={last_res.returncode}")
                if last_res.stdout:
                    diag.append('\n--- stdout ---\n' + last_res.stdout[-2000:])
                if last_res.stderr:
                    diag.append('\n--- stderr ---\n' + last_res.stderr[-2000:])

            log_file = tmp / 'report.log'
            if log_file.exists():
                try:
                    log_text = log_file.read_text(encoding='utf-8', errors='ignore')
                    diag.append('\n--- report.log (tail) ---\n' + log_text[-3000:])
                except Exception:
                    pass

            raise Exception('LaTeX compilation failed; no PDF produced.\n' + '\n'.join(diag))
    except ImportError:
        raise Exception('Jinja2 required for PDF generation')
    except FileNotFoundError:
        raise Exception('LaTeX compiler (pdflatex/xelatex/lualatex) not found on PATH — install TeX Live / MiKTeX or add to PATH')
    except Exception as e:
        raise Exception(f'PDF generation failed: {e}')