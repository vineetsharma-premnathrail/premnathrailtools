import io
import math
import os
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any as TypingAny, Any, Dict, List, Optional

# Shared templates folder: backend/app/utils/templates/
_TEMPLATES_DIR = Path(__file__).resolve().parents[5] / "utils" / "templates"


def generate_braking_pdf_report(context: Dict[str, Any]) -> io.BytesIO:
    """Generate PDF report using LaTeX template + Jinja2."""
    try:
        from jinja2 import Environment, FileSystemLoader
        import shutil

        env = Environment(loader=FileSystemLoader(str(_TEMPLATES_DIR)))
        template = env.get_template('template.tex')

        try:
            latex_content = template.render(**context)
        except Exception as exc:
            raise Exception(f"Jinja2 template rendering failed: {exc}")

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            tex_file = temp_path / 'report.tex'
            tex_file.write_text(latex_content, encoding='utf-8')

            for img_name in ["logo.JPG", "logo-1.JPG", "breaking distance table.png"]:
                img_src = _TEMPLATES_DIR / img_name
                if img_src.exists():
                    shutil.copy(img_src, temp_path / img_name)

            pdf_path = temp_path / 'report.pdf'
            last_res = None
            # openin_any=p restricts \input/\include file resolution to the
            # temp working directory — defense in depth alongside escaping
            # user-supplied fields before they reach the template.
            tex_env = {**os.environ, "openin_any": "p"}

            for compiler in ['pdflatex', 'xelatex', 'lualatex']:
                try:
                    last_res = subprocess.run(
                        [compiler, '-interaction=nonstopmode', '-no-shell-escape', 'report.tex'],
                        cwd=temp_dir, capture_output=True, text=True, timeout=30, env=tex_env,
                    )
                except FileNotFoundError:
                    last_res = None
                    continue
                if pdf_path.exists():
                    return io.BytesIO(pdf_path.read_bytes())

            diag = []
            if last_res:
                diag.append(f"returncode={last_res.returncode}")
                diag.append(last_res.stdout or '')
                diag.append(last_res.stderr or '')
            log_file = temp_path / 'report.log'
            if log_file.exists():
                try:
                    diag.append(log_file.read_text(encoding='utf-8', errors='ignore')[-2000:])
                except Exception:
                    pass
            raise Exception('LaTeX compilation failed; no PDF produced.\n' + '\n'.join(diag))

    except FileNotFoundError:
        raise Exception("LaTeX compiler not found. Install texlive or MiKTeX.")
    except ImportError:
        raise Exception("Jinja2 not installed. Run: pip install Jinja2")
    except Exception as e:
        raise Exception(f"PDF generation failed: {e}")


def create_braking_docx_report(
    results_table_rows: List[Any],
    context: Dict[str, Any],
) -> io.BytesIO:
    """Detailed DOCX report with step-by-step formulas for Rail and Rail+Road modes."""
    from docx import Document

    G = 9.81
    doc = Document()
    doc.add_heading("Braking Performance Calculation Report", 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    if context.get('doc_no'):
        doc.add_paragraph(f"Document No: {context['doc_no']}")
    if context.get('made_by'):
        doc.add_paragraph(
            f"Prepared by: {context['made_by']}  |  "
            f"Checked: {context.get('checked_by', '')}  |  "
            f"Approved: {context.get('approved_by', '')}"
        )

    # 1. INPUTS
    doc.add_heading("1. Inputs", level=1)
    mass_kg = float(context.get('mass_kg', 0))
    weight_n = float(context.get('weight_n', mass_kg * G))
    reaction_time = float(context.get('reaction_time', 0))
    num_wheels = int(context.get('number_of_wheels', 0))
    wheel_dia = float(context.get('wheel_dia', 0))
    mu = float(context.get('friction_coefficient', 0.7))
    grad_input = float(context.get('gradient_input', 0))
    grad_type = context.get('gradient_type', '')
    max_braking_f = float(context.get('max_braking_force', 0))

    p = doc.add_paragraph()
    p.add_run(f"  Vehicle Mass:           {mass_kg} kg\n")
    p.add_run(f"  Weight (N):             {weight_n:.2f} N  (= {mass_kg} × {G})\n")
    p.add_run(f"  Reaction Time:          {reaction_time} s\n")
    p.add_run(f"  Number of Wheels:       {num_wheels}\n")
    p.add_run(f"  Wheel Diameter:         {wheel_dia} mm\n")
    p.add_run(f"  Rail Gradient:          {grad_input} ({grad_type})\n")
    p.add_run(f"  Friction Coefficient μ: {mu}\n")

    # 2. FORCE CAPABILITY
    doc.add_heading("2. Step 1: Force Capability from Standard Speed–Distance Data", level=1)
    p = doc.add_paragraph()
    p.add_run("Formula: a_required = v² / (2 × d)  →  F_required = m × a_required\n")
    p.add_run("(Standard BRAKING_DATA per DIN EN 15746-2)\n\n")

    old_data: Dict = context.get('old_data_for_report', {})
    if old_data:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, h in enumerate(['Speed (km/h)', 'v (m/s)', 'Std Dist (m)', 'Decel (m/s²)', 'Reaction Dist (m)', 'Brake Force (N)']):
            hdr[i].text = h
        for spd, row_d in sorted(old_data.items()):
            rc = table.add_row().cells
            rc[0].text = str(spd)
            rc[1].text = str(row_d.get('speed_ms', ''))
            rc[2].text = str(row_d.get('braking_distance', ''))
            rc[3].text = str(row_d.get('deceleration', ''))
            rc[4].text = str(row_d.get('reaction_distance', ''))
            rc[5].text = str(row_d.get('braking_force', ''))

    p2 = doc.add_paragraph()
    r = p2.add_run(f"\n  Max Braking Force selected = {max_braking_f:.2f} N\n")
    r.bold = True

    # 3. SCENARIO FORCES
    doc.add_heading("3. Step 2: Scenario Net Forces (Straight / Up / Down)", level=1)
    p = doc.add_paragraph()
    p.add_run("  grav_force = weight_N × sin(θ°)\n")
    p.add_run("  Straight Track:  F_net = F_max\n")
    p.add_run("  Moving Up:       F_net = F_max + grav_force\n")
    p.add_run("  Moving Down:     F_net = F_max − grav_force\n\n")

    rail_calcs: List[Dict] = context.get('rail_detailed_calcs', [])
    shown: set = set()
    for rc in rail_calcs:
        key = (rc.get('scenario'), rc.get('gradient_value'))
        if key in shown:
            continue
        shown.add(key)
        ang = rc.get('angle_deg', 0)
        grav_f = rc.get('f_g', 0)
        f_net = rc.get('f_net', 0)
        s = rc.get('scenario')
        grad_v = rc.get('gradient_value', 0)
        p.add_run(f"  Gradient {grad_v} ({grad_type}) → θ = {ang}°,  sin(θ) = {math.sin(math.radians(ang)):.6f}\n")
        p.add_run(f"  grav_force = {weight_n:.2f} × {math.sin(math.radians(ang)):.6f} = {grav_f:.2f} N\n")
        if s == "Straight Track":
            p.add_run(f"  [{s}]  F_net = {max_braking_f:.2f} N\n\n")
        elif s == "Moving up":
            p.add_run(f"  [{s}]  F_net = {max_braking_f:.2f} + {grav_f:.2f} = {f_net:.2f} N\n\n")
        elif s == "Moving down":
            p.add_run(f"  [{s}]  F_net = {max_braking_f:.2f} − {grav_f:.2f} = {f_net:.2f} N\n\n")

    # 4. DECELERATION & DISTANCES
    doc.add_heading("4. Step 3: Deceleration & Stopping Distances", level=1)
    p = doc.add_paragraph()
    p.add_run("  a = |F_net / m|\n")
    p.add_run("  braking_dist  = v² / (2 × a)\n")
    p.add_run("  reaction_dist = v × t_reaction\n")
    p.add_run("  total_dist    = reaction_dist + braking_dist\n\n")

    if rail_calcs:
        table2 = doc.add_table(rows=1, cols=7)
        table2.style = 'Table Grid'
        h = table2.rows[0].cells
        for i, hd in enumerate(['Scenario', 'Speed (km/h)', 'Gradient', 'Decel (m/s²)', 'React Dist (m)', 'Brake Dist (m)', 'Total Dist (m)']):
            h[i].text = hd
        for rc in rail_calcs:
            r2 = table2.add_row().cells
            r2[0].text = str(rc.get('scenario', ''))
            r2[1].text = str(rc.get('speed_kmh', ''))
            r2[2].text = str(rc.get('gradient_value', ''))
            r2[3].text = str(rc.get('a_deceleration', ''))
            r2[4].text = str(rc.get('reaction_distance', ''))
            r2[5].text = str(rc.get('braking_distance', ''))
            r2[6].text = str(rc.get('total_stopping_distance', ''))

        ex = rail_calcs[0]
        doc.add_heading("Worked Example (first entry)", level=2)
        p = doc.add_paragraph()
        ex_v = ex.get('v_ms', 0)
        ex_fn = ex.get('f_net', 0)
        ex_a = ex.get('a_deceleration', 0)
        ex_rd = ex.get('reaction_distance', 0)
        ex_bd = ex.get('braking_distance', 0)
        ex_td = ex.get('total_stopping_distance', 0)
        p.add_run(f"  Speed = {ex.get('speed_kmh')} km/h = {ex_v} m/s,  F_net = {ex_fn} N,  m = {mass_kg} kg\n")
        p.add_run(f"  a = |{ex_fn} / {mass_kg}| = {ex_a} m/s²\n")
        p.add_run(f"  braking_dist  = {ex_v}² / (2 × {ex_a}) = {ex.get('v_ms_squared', ex_v**2):.2f} / {2*ex_a:.4f} = {ex_bd} m\n")
        p.add_run(f"  reaction_dist = {ex_v} × {reaction_time} = {ex_rd} m\n")
        p.add_run(f"  total_dist    = {ex_rd} + {ex_bd} = {ex_td} m\n").bold = True

    # 5. GBR
    doc.add_heading("5. Step 4: Gross Braking Ratio (GBR)", level=1)
    gbr = float(context.get('gbr', 0))
    p = doc.add_paragraph()
    p.add_run("Formula: GBR = (F_max / (m × g)) × 100\n")
    p.add_run(f"  = ({max_braking_f:.2f} / ({mass_kg} × {G})) × 100\n")
    p.add_run(f"  GBR = {gbr:.2f} %\n").bold = True

    # 6. ROAD MODE
    road_calcs: List[Dict] = context.get('road_detailed_calcs', [])
    if road_calcs:
        doc.add_heading("6. Step 5: Road Mode Calculations", level=1)
        p = doc.add_paragraph()
        p.add_run(f"  μ = {mu}\n")
        p.add_run("  normal   = weight_N × cos(θ°)\n")
        p.add_run("  friction = μ × normal\n")
        p.add_run("  Straight: F_net = friction\n")
        p.add_run("  Up:       F_net = friction + grav_force\n")
        p.add_run("  Down:     F_net = friction − grav_force\n\n")

        table3 = doc.add_table(rows=1, cols=7)
        table3.style = 'Table Grid'
        h3 = table3.rows[0].cells
        for i, hd in enumerate(['Scenario', 'Speed (km/h)', 'Gradient', 'Normal (N)', 'Friction (N)', 'F_net (N)', 'Total Dist (m)']):
            h3[i].text = hd
        for rc in road_calcs:
            rr = table3.add_row().cells
            rr[0].text = str(rc.get('scenario', ''))
            rr[1].text = str(rc.get('speed_kmh', ''))
            rr[2].text = str(rc.get('gradient_value', ''))
            rr[3].text = str(rc.get('normal_force', ''))
            rr[4].text = str(rc.get('fb_friction', ''))
            rr[5].text = str(rc.get('f_net', ''))
            rr[6].text = str(rc.get('total_stopping_distance', ''))

    # 7. RESULTS SUMMARY
    doc.add_heading("7. Results Summary", level=1)
    if results_table_rows:
        table4 = doc.add_table(rows=1, cols=7)
        table4.style = 'Table Grid'
        h4 = table4.rows[0].cells
        for i, hd in enumerate(['Mode', 'Scenario', 'Speed (km/h)', 'Gradient', 'Decel (m/s²)', 'Total Dist (m)', 'Status']):
            h4[i].text = hd
        for row in results_table_rows:
            rc4 = table4.add_row().cells
            rc4[0].text = str(row.get('mode', ''))
            rc4[1].text = str(row.get('scenario', ''))
            rc4[2].text = str(row.get('speed', ''))
            rc4[3].text = str(row.get('gradient', ''))
            rc4[4].text = str(row.get('decel', ''))
            rc4[5].text = str(row.get('total', ''))
            rc4[6].text = str(row.get('status', ''))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
