"""DOCX builder for the CRM "Technical Offer Request" — the document emailed
to R&D when a Technical Offer Request is raised from an Inquiry or Tender.

Reuses the same cropped letterhead images as quotation_pdf.py (logo.JPG /
logo-1.JPG), stamped full page-width edge-to-edge in the header/footer via a
floating picture anchored relative to the page (not the column/margins), the
same technique mom_docx.py uses for its logo — just anchored to "page" instead
of "column" so the banner bleeds to both page edges like the PDF does.
"""

from __future__ import annotations

import io
import os
from typing import Any, Mapping

try:
    from docx import Document
    from docx.shared import Pt, Inches, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    Document = None

try:
    from PIL import Image as PILImage, ImageChops
except ImportError:
    PILImage = None

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "..", "utils", "templates")
HEADER_LOGO_PATH = os.path.join(TEMPLATES_DIR, "logo.JPG")
FOOTER_LOGO_PATH = os.path.join(TEMPLATES_DIR, "logo-1.JPG")

FONT_NAME = "Calibri"
ACCENT = RGBColor(0x1C, 0x21, 0x3E)
MUTED = RGBColor(0x57, 0x53, 0x4E)

PAGE_WIDTH_EMU = Emu(Inches(8.27)) if Document else 0  # A4 width


def _cropped_content_bytes(path: str) -> bytes | None:
    """Same trim-to-content-bbox treatment as quotation_pdf._load_content_cropped —
    the source JPEGs carry asymmetric white padding, so stretching the raw file to
    full page width leaves a blank strip on one side."""
    if PILImage is None or not os.path.exists(path):
        return None
    try:
        img = PILImage.open(path).convert("RGB")
        bg = PILImage.new("RGB", img.size, (255, 255, 255))
        diff = ImageChops.difference(img, bg)
        bbox = diff.getbbox()
        if bbox:
            img = img.crop(bbox)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return None


def _image_aspect(image_bytes: bytes) -> float:
    """height / width, so a target width can preserve the source aspect ratio."""
    try:
        img = PILImage.open(io.BytesIO(image_bytes))
        w, h = img.size
        return h / w
    except Exception:
        return 0.12


def _add_full_width_banner(paragraph: Any, image_bytes: bytes, offset_y_emu: int) -> None:
    """Insert `image_bytes` as a floating picture spanning the full page width,
    anchored relative to the page (not the margins) so it bleeds edge-to-edge."""
    aspect = _image_aspect(image_bytes)
    width_emu = int(PAGE_WIDTH_EMU)
    height_emu = int(width_emu * aspect)

    stream = io.BytesIO(image_bytes)
    run = paragraph.add_run()
    run.add_picture(stream, width=Emu(width_emu), height=Emu(height_emu))

    drawing = run._element.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))
    extent = inline.find(qn("wp:extent"))
    doc_pr = inline.find(qn("wp:docPr"))
    graphic = inline.find(qn("a:graphic"))
    cnv_graphic_frame_pr = inline.find(qn("wp:cNvGraphicFramePr"))

    anchor = OxmlElement("wp:anchor")
    for attr, val in (
        ("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
        ("allowOverlap", "1"), ("layoutInCell", "1"), ("locked", "0"),
        ("behindDoc", "0"), ("simplePos", "0"), ("relativeHeight", "251658240"),
    ):
        anchor.set(attr, val)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "page")
    offset_h = OxmlElement("wp:posOffset")
    offset_h.text = "0"
    position_h.append(offset_h)
    anchor.append(position_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "page")
    offset_v = OxmlElement("wp:posOffset")
    offset_v.text = str(offset_y_emu)
    position_v.append(offset_v)
    anchor.append(position_v)

    anchor.append(extent)

    effect_extent = OxmlElement("wp:effectExtent")
    for attr in ("l", "t", "r", "b"):
        effect_extent.set(attr, "0")
    anchor.append(effect_extent)

    anchor.append(OxmlElement("wp:wrapNone"))
    anchor.append(doc_pr)
    anchor.append(cnv_graphic_frame_pr)
    anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)


def _heading(doc: Any, number: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{number}. {text}")
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12.5)
    run.font.color.rgb = ACCENT
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1C213E")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)


def _paragraph(doc: Any, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(10.5)


def _bullet(doc: Any, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(10.5)


def _spec_table(doc: Any, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, label in enumerate(("Field", "Value")):
        hdr[i].text = label
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(10)


def build_technical_offer_docx(ctx: Mapping[str, Any]) -> io.BytesIO:
    """Build the numbered-section "Technical Offer" .docx from a plain dict context.

    Expected keys: offer_number, universal_id, org_name, org_type, org_address, org_gst_number,
    org_city, org_state, contact_name, contact_designation, contact_department, contact_mobile,
    contact_email, product_category, product, quantity_display, required_delivery_date,
    delivery_location, inspection_req, warranty_req, product_spec, requirement_desc,
    detailed_requirement, project_details, documents_link, raised_by, raised_at. All are
    optional except offer_number/universal_id — every section degrades gracefully to a
    placeholder ("Not specified."/"Not provided.") rather than a fabricated value when its
    underlying field is empty on the record.
    """
    if Document is None:
        raise Exception("python-docx library is not installed. Please run: pip install python-docx")

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(1.05)
    section.bottom_margin = Inches(0.75)

    header_bytes = _cropped_content_bytes(HEADER_LOGO_PATH)
    footer_bytes = _cropped_content_bytes(FOOTER_LOGO_PATH)
    if header_bytes:
        header_para = section.header.paragraphs[0]
        _add_full_width_banner(header_para, header_bytes, offset_y_emu=0)
    if footer_bytes:
        footer_para = section.footer.paragraphs[0]
        aspect = _image_aspect(footer_bytes)
        footer_height_emu = int(int(PAGE_WIDTH_EMU) * aspect)
        page_height_emu = int(Emu(Inches(11.69)))
        _add_full_width_banner(footer_para, footer_bytes, offset_y_emu=page_height_emu - footer_height_emu)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("TECHNICAL OFFER")
    title_run.bold = True
    title_run.font.name = FONT_NAME
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = ACCENT

    product_line = doc.add_paragraph()
    product_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    product_line_run = product_line.add_run(ctx.get("product") or ctx.get("product_category") or ctx.get("universal_id") or "-")
    product_line_run.bold = True
    product_line_run.font.name = FONT_NAME
    product_line_run.font.size = Pt(13)

    for_line = doc.add_paragraph()
    for_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for_line_run = for_line.add_run(f"For {ctx.get('org_name') or '-'}")
    for_line_run.font.name = FONT_NAME
    for_line_run.font.size = Pt(11)
    for_line_run.font.color.rgb = MUTED

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    subtitle_run = subtitle.add_run(f"Offer No. {ctx.get('offer_number', '-')}  ·  Ref. {ctx.get('universal_id', '-')}")
    subtitle_run.font.name = FONT_NAME
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = MUTED

    # 1. Customer / Organization Details
    _heading(doc, 1, "Customer / Organization Details")
    org_rows = [("Organization", ctx.get("org_name") or "-")]
    if ctx.get("org_gst_number"):
        org_rows.append(("GST Number", ctx["org_gst_number"]))
    org_rows.append(("Organization Type", ctx.get("org_type") or "-"))
    city_state = ", ".join(v for v in (ctx.get("org_city"), ctx.get("org_state")) if v)
    if city_state:
        org_rows.append(("City / State", city_state))
    if ctx.get("org_address"):
        org_rows.append(("Address", ctx["org_address"]))
    _spec_table(doc, org_rows)

    # 2. Contact Person
    if any(ctx.get(k) for k in ("contact_name", "contact_designation", "contact_department", "contact_mobile", "contact_email")):
        _heading(doc, 2, "Contact Person")
        contact_rows = []
        for label, key in (
            ("Name", "contact_name"), ("Designation", "contact_designation"),
            ("Department", "contact_department"), ("Mobile", "contact_mobile"), ("Email", "contact_email"),
        ):
            if ctx.get(key):
                contact_rows.append((label, ctx[key]))
        _spec_table(doc, contact_rows)

    # 3. Technical Requirement Summary
    _heading(doc, 3, "Technical Requirement Summary")
    req_rows = []
    for label, key in (
        ("Category", "product_category"), ("Product", "product"), ("Quantity", "quantity_display"),
        ("Required Delivery Date", "required_delivery_date"), ("Delivery Location", "delivery_location"),
        ("Inspection", "inspection_req"), ("Warranty", "warranty_req"),
    ):
        if ctx.get(key):
            req_rows.append((label, str(ctx[key])))
    if req_rows:
        _spec_table(doc, req_rows)
    else:
        _paragraph(doc, "No structured requirement fields have been filled in for this record yet.")

    # 4. Technical Specification
    _heading(doc, 4, "Technical Specification")
    _paragraph(doc, ctx.get("product_spec") or "Not specified.")

    # 5. Requirement Summary
    _heading(doc, 5, "Requirement Summary")
    _paragraph(doc, ctx.get("requirement_desc") or ctx.get("detailed_requirement") or "Not provided.")

    # 6. Project Background
    _heading(doc, 6, "Project Background")
    _paragraph(doc, ctx.get("project_details") or ctx.get("detailed_requirement") or "Not provided.")

    # 7. Inspection, Delivery and Warranty Requirements
    _heading(doc, 7, "Inspection, Delivery and Warranty Requirements")
    any_bullet = False
    if ctx.get("quantity_display"):
        _bullet(doc, f"Supply quantity: {ctx['quantity_display']}.")
        any_bullet = True
    if ctx.get("required_delivery_date"):
        _bullet(doc, f"Delivery deadline: {ctx['required_delivery_date']}.")
        any_bullet = True
    if ctx.get("delivery_location"):
        _bullet(doc, f"Delivery location: {ctx['delivery_location']}.")
        any_bullet = True
    if ctx.get("inspection_req"):
        _bullet(doc, ctx["inspection_req"])
        any_bullet = True
    if ctx.get("warranty_req"):
        _bullet(doc, f"Warranty: {ctx['warranty_req']}.")
        any_bullet = True
    if not any_bullet:
        _paragraph(doc, "No inspection, delivery, or warranty requirements have been specified yet.")

    # 8. Technical Compliance Statement
    _heading(doc, 8, "Technical Compliance Statement")
    _paragraph(
        doc,
        "The proposed technical offer shall be evaluated against the above customer requirements. "
        "All deviations, if any, shall be explicitly identified in the final technical compliance statement.",
    )

    # 9. Document Control
    _heading(doc, 9, "Document Control")
    _spec_table(doc, [
        ("Document Title", f"Technical Offer – {ctx.get('product') or ctx.get('product_category') or ctx.get('universal_id') or '-'}"),
        ("Customer", ctx.get("org_name") or "-"),
        ("Revision", "00"),
        ("Issue Date", ctx.get("raised_at", "-").split(",")[0] if ctx.get("raised_at") else "-"),
        ("Status", "Technical Offer Request"),
    ])

    if ctx.get("documents_link"):
        link_p = doc.add_paragraph()
        link_p.paragraph_format.space_before = Pt(10)
        link_label = link_p.add_run("Reference documents already on file: ")
        link_label.bold = True
        link_label.font.name = FONT_NAME
        link_label.font.size = Pt(9.5)
        link_run = link_p.add_run(ctx["documents_link"])
        link_run.font.name = FONT_NAME
        link_run.font.size = Pt(9.5)
        link_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        link_run.font.underline = True

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note_run = note.add_run(
        f"Please quote Technical Offer No. {ctx.get('offer_number', '-')} in your technical offer for tracking."
    )
    note_run.bold = True
    note_run.font.name = FONT_NAME
    note_run.font.size = Pt(10.5)

    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_p.paragraph_format.space_before = Pt(20)
    end_run = end_p.add_run("End of Technical Offer")
    end_run.italic = True
    end_run.font.name = FONT_NAME
    end_run.font.size = Pt(9)
    end_run.font.color.rgb = MUTED

    footer_note = doc.add_paragraph()
    footer_note.paragraph_format.space_before = Pt(8)
    footer_note_run = footer_note.add_run(
        f"Raised by {ctx.get('raised_by', '-')} on {ctx.get('raised_at', '-')}."
    )
    footer_note_run.font.name = FONT_NAME
    footer_note_run.font.size = Pt(9)
    footer_note_run.font.color.rgb = MUTED

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
