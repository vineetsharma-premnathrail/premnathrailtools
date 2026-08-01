"""DOCX builder for the CRM "Minutes of Meeting" export.

Ported directly from a field-by-field recreation of a real, signed-off
F/HRD/4 MOM document (single merged-cell table, landscape page, Premnath
Rail letterhead, floating logo, Calibri throughout, zero paragraph
spacing) — the structure and low-level OOXML helpers below mirror that
recreation exactly; only the text content is parameterized per inquiry.
"""

from __future__ import annotations

import io
import os
from typing import Any, Mapping, Sequence

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    Document = None

LOGO_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "..", "utils", "templates", "premnath_logo_mark.png")
FONT_NAME = "Calibri"

# Column widths in twips — lifted directly from the source template's
# <w:tblGrid>, which Word treats as authoritative for a fixed-layout table
# (per-cell width alone gets overridden back to equal columns otherwise).
COL_WIDTHS_TWIPS = [688, 6617, 2801, 2201, 1282]


# ---------------------------------------------------------------------------
# low level helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell: Any, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_cell_vertical_alignment(cell: Any, align: str = "center") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), align)
    tc_pr.append(v_align)


def _add_run(paragraph: Any, text: str, bold: bool = False, size: float = 10, color: str | None = None) -> Any:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _zero_spacing(paragraph: Any) -> Any:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    return paragraph


def _clear_paragraph(cell: Any) -> Any:
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    return _zero_spacing(paragraph)


def _new_paragraph(cell: Any) -> Any:
    return _zero_spacing(cell.add_paragraph())


def _add_floating_picture(
    paragraph: Any, image_path: str, width_emu: int, height_emu: int,
    offset_x_emu: int, offset_y_emu: int, rel_height: int = 487490048,
) -> Any:
    """Insert `image_path` as a floating (anchored) picture that sits behind
    the text at a fixed position — like the logo in the source letterhead —
    instead of a normal inline image that flows with, and gets pushed around
    by, the surrounding paragraphs."""
    run = paragraph.add_run()
    run.add_picture(image_path, width=Emu(width_emu), height=Emu(height_emu))

    drawing = run._element.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))

    extent = inline.find(qn("wp:extent"))
    doc_pr = inline.find(qn("wp:docPr"))
    graphic = inline.find(qn("a:graphic"))
    cnv_graphic_frame_pr = inline.find(qn("wp:cNvGraphicFramePr"))

    anchor = OxmlElement("wp:anchor")
    for attr, val in (
        ("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
        ("allowOverlap", "1"), ("layoutInCell", "1"), ("locked", "0"),
        ("behindDoc", "1"), ("simplePos", "0"), ("relativeHeight", str(rel_height)),
    ):
        anchor.set(attr, val)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "column")
    offset_h = OxmlElement("wp:posOffset")
    offset_h.text = str(offset_x_emu)
    position_h.append(offset_h)
    anchor.append(position_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "paragraph")
    offset_v = OxmlElement("wp:posOffset")
    offset_v.text = str(offset_y_emu)
    position_v.append(offset_v)
    anchor.append(position_v)

    anchor.append(extent)

    effect_extent = OxmlElement("wp:effectExtent")
    for attr in ("l", "t", "r", "b"):
        effect_extent.set(attr, "0")
    anchor.append(effect_extent)

    anchor.append(OxmlElement("wp:wrapNone"))
    anchor.append(doc_pr)
    anchor.append(cnv_graphic_frame_pr)
    anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)
    return run


def _set_col_widths(table: Any, widths_twips: Sequence[int]) -> None:
    """Force fixed column widths on every row and on the table's own
    <w:tblGrid> — Word renders merged-cell tables off the grid definition,
    so per-cell width alone isn't enough."""
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Twips(widths_twips[idx])
    tbl = table._tbl
    tbl_grid = tbl.find(qn("w:tblGrid"))
    for grid_col, width in zip(tbl_grid.findall(qn("w:gridCol")), widths_twips):
        grid_col.set(qn("w:w"), str(width))


def _attendee_lines(members: Sequence[Mapping[str, Any]]) -> list[str]:
    if not members:
        return ["-"]
    lines = []
    for member in members:
        label = member.get("name") or "-"
        if member.get("designation"):
            label = f"{label} ({member['designation']})"
        lines.append(label)
    return lines


# ---------------------------------------------------------------------------
# build document
# ---------------------------------------------------------------------------

def build_mom_docx(ctx: Mapping[str, Any]) -> io.BytesIO:
    """Build the MOM .docx from a plain-dict context.

    Expected keys: org_name, subject, meeting_date (already formatted str),
    pew_members / client_members (list of {name, designation}), and
    activities (list of {observation, action_plan, responsibility, target}).
    """
    if Document is None:
        raise Exception("python-docx library is not installed. Please run: pip install python-docx")

    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Twips(1380)
    section.bottom_margin = Twips(280)
    section.left_margin = Twips(720)
    section.right_margin = Twips(1080)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    activities = ctx.get("activities") or []
    table = doc.add_table(rows=10 + len(activities), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge, sz in (("top", 18), ("left", 18), ("bottom", 18), ("right", 18), ("insideH", 8), ("insideV", 8)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)

    rows = table.rows

    # ROWS 0-3: letterhead block — logo + company name, cols 0-2 merged
    # horizontally on each row, then col 0 merged vertically across the four.
    for r in range(4):
        rows[r].cells[0].merge(rows[r].cells[2])
    head_cell = rows[0].cells[0]
    for r in range(1, 4):
        head_cell = head_cell.merge(rows[r].cells[0])

    p = _clear_paragraph(head_cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14.5)
    if os.path.exists(LOGO_PATH):
        try:
            _add_floating_picture(p, LOGO_PATH, width_emu=1487805, height_emu=576580, offset_x_emu=74676, offset_y_emu=65377)
        except Exception:
            pass
    _add_run(p, "Premnath Engineering Works", bold=True, size=14.5)
    _set_cell_vertical_alignment(head_cell, "center")

    for i, (label, value) in enumerate([
        ("Format No :", "F/HRD/4"),
        ("Effective Date :", "01.02.2025"),
        ("Rev. No :", "00"),
        ("Rev Date :", "00"),
    ]):
        _add_run(_clear_paragraph(rows[i].cells[3]), label, bold=True, size=11)
        _add_run(_clear_paragraph(rows[i].cells[4]), value, bold=True, size=11)

    # ROW 4: black "MINUTES OF MEETING" banner.
    banner_cell = rows[4].cells[0]
    for c in range(1, 5):
        banner_cell = banner_cell.merge(rows[4].cells[c])
    _shade_cell(banner_cell, "000000")
    p = _clear_paragraph(banner_cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "MINUTES OF MEETING", bold=True, size=18, color="FFFFFF")

    # ROW 5: thin spacer.
    spacer_cell = rows[5].cells[0]
    for c in range(1, 5):
        spacer_cell = spacer_cell.merge(rows[5].cells[c])
    _clear_paragraph(spacer_cell)

    # ROW 6: client name.
    client_cell = rows[6].cells[0]
    for c in range(1, 5):
        client_cell = client_cell.merge(rows[6].cells[c])
    _add_run(_clear_paragraph(client_cell), f"CLIENT: {ctx.get('org_name') or '-'}", bold=True, size=14.5)

    # ROW 7: subject / date.
    subject_cell = rows[7].cells[0].merge(rows[7].cells[1])
    date_cell = rows[7].cells[2].merge(rows[7].cells[3]).merge(rows[7].cells[4])
    _add_run(_clear_paragraph(subject_cell), f"SUBJECT : {ctx.get('subject') or '-'}", bold=True, size=14.5)
    _add_run(_clear_paragraph(date_cell), f"DATE: {ctx.get('meeting_date') or '-'}", bold=True, size=13)

    # ROW 8: PEW / client members present.
    pew_cell = rows[8].cells[0].merge(rows[8].cells[1])
    client_members_cell = rows[8].cells[2].merge(rows[8].cells[3]).merge(rows[8].cells[4])

    _add_run(_clear_paragraph(pew_cell), "PEW MEMBER PRESENT:-", bold=True, size=13)
    for line in _attendee_lines(ctx.get("pew_members") or []):
        _add_run(_new_paragraph(pew_cell), line, bold=False, size=13)

    _add_run(_clear_paragraph(client_members_cell), "CLIENT MEMBERS PRESENT -", bold=True, size=13)
    for line in _attendee_lines(ctx.get("client_members") or []):
        _add_run(_new_paragraph(client_members_cell), line, bold=False, size=13)

    # ROW 9: column headers.
    headers = ["S NO", "OBSERVATION", "ACTION PLAN", "RESPONSIBILITY", "TARGET\nDATE"]
    for c, text in enumerate(headers):
        cell = rows[9].cells[c]
        lines = text.split("\n")
        p = _clear_paragraph(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, lines[0], bold=True, size=13)
        for extra in lines[1:]:
            p2 = _new_paragraph(cell)
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p2, extra, bold=True, size=13)
        _set_cell_vertical_alignment(cell, "center")

    # ROWS 10+: one per included activity.
    for i, activity in enumerate(activities):
        row_idx = 10 + i
        cells = rows[row_idx].cells

        p = _clear_paragraph(cells[0])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, str(i + 1), bold=False, size=11)

        _add_run(_clear_paragraph(cells[1]), activity.get("observation") or "-", bold=False, size=10)
        _add_run(_clear_paragraph(cells[2]), activity.get("action_plan") or "-", bold=False, size=10)

        p = _clear_paragraph(cells[3])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, activity.get("responsibility") or "-", bold=False, size=10)

        p = _clear_paragraph(cells[4])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, activity.get("target") or "-", bold=False, size=10)

    _set_col_widths(table, COL_WIDTHS_TWIPS)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
